"""
Minuta Word (tono dirección): inventario de la carpeta compartida Actuarial frente a la
Solicitud consolidada Definitivo v2 — Mesa Actuarial.

El **detalle de reaseguro, TPA en salud y PAS (SAA-04-1443)** no se desarrolla aquí: vive en el
informe confidencial generado por `scripts/generar_informe_confidencial_reaseguro_tpa_salud_docx.py`.

Logo: `plantillas/logo-la-internacional.png`. Encabezado en tabla centrada en la página, con
márgenes laterales implícitos al no usar todo el ancho útil.
"""
from __future__ import annotations

import argparse
import io
import re
import tempfile
import unicodedata
from datetime import date, datetime
from pathlib import Path

from docx import Document
from docx.enum.table import WD_ALIGN_VERTICAL, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor

TITLE_BLUE = RGBColor(0x1B, 0x3A, 0x5C)
ACCENT_GOLD = RGBColor(0xC9, 0xA2, 0x27)

ROOT_REPORTERIA = Path(__file__).resolve().parents[1]
WORKSPACE = ROOT_REPORTERIA.parent
DEFAULT_LOGO_SVG = WORKSPACE / "La-Internacional-BI" / "public" / "Images" / "Logo horizontal.svg"
BUNDLED_LOGO_PNG = ROOT_REPORTERIA / "plantillas" / "logo-la-internacional.png"
DEFAULT_OUT = (
    WORKSPACE
    / "Reportes"
    / "Mesa Actuarial"
    / "Minutas Reuniones"
    / "MIN-POR-ACT-20260424-Mesa-Actuarial-vs-Solicitud_actualizado-2026-04-24.docx"
)


def _set_run_font(run, *, bold: bool = False, italic: bool = False, size: int = 11, color=None) -> None:
    run.bold = bold
    run.italic = italic
    run.font.name = "Arial"
    run.font.size = Pt(size)
    run._element.rPr.rFonts.set(qn("w:eastAsia"), "Arial")
    if color is not None:
        run.font.color.rgb = color


def _heading(doc: Document, text: str, level: int = 1) -> None:
    h = doc.add_heading(text, level=level)
    # Jerarquía por debajo del titular de portada (18 pt): secciones 14 / 11 pt.
    sec_size = 14 if level == 1 else 11
    for run in h.runs:
        _set_run_font(run, bold=True, size=sec_size, color=TITLE_BLUE)
    h.paragraph_format.space_before = Pt(12)
    h.paragraph_format.space_after = Pt(6)


def _p(doc: Document, text: str, *, bold: bool = False, italic: bool = False, size: int = 11) -> None:
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(6)
    r = p.add_run(text)
    _set_run_font(r, bold=bold, italic=italic, size=size)


def _add_markdown_runs_to_paragraph(
    p,
    text: str,
    *,
    size: int = 11,
    italic_default: bool = False,
) -> None:
    """Añade runs a un párrafo existente interpretando **negritas** (no atraviesa saltos de línea)."""
    parts = re.split(r"(\*\*.+?\*\*)", text)
    for part in parts:
        if not part:
            continue
        if part.startswith("**") and part.endswith("**") and len(part) > 4:
            inner = part[2:-2].replace("\n", " ")
            r = p.add_run(inner)
            _set_run_font(r, bold=True, size=size)
        else:
            r = p.add_run(part)
            _set_run_font(r, italic=italic_default, size=size)


def _p_rich(doc: Document, text: str, *, size: int = 11, italic_default: bool = False) -> None:
    """Párrafo con **negritas** estilo Markdown (consistente en informes CONF)."""
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(6)
    _add_markdown_runs_to_paragraph(p, text, size=size, italic_default=italic_default)


def _bullet_rich(doc: Document, text: str, *, size: int = 11) -> None:
    p = doc.add_paragraph(style="List Bullet")
    p.paragraph_format.space_after = Pt(4)
    _add_markdown_runs_to_paragraph(p, text, size=size, italic_default=False)


def _add_table(doc: Document, headers: list[str], rows: list[list[str]]) -> None:
    table = doc.add_table(rows=1 + len(rows), cols=len(headers))
    table.style = "Table Grid"
    for j, h in enumerate(headers):
        cell = table.rows[0].cells[j]
        cell.text = ""
        p = cell.paragraphs[0]
        run = p.add_run(h)
        _set_run_font(run, bold=True, size=10)
    for i, row in enumerate(rows, start=1):
        for j, val in enumerate(row):
            cell = table.rows[i].cells[j]
            cell.text = ""
            p = cell.paragraphs[0]
            sval = str(val)
            if "**" in sval:
                _add_markdown_runs_to_paragraph(p, sval, size=10)
            else:
                run = p.add_run(sval)
                _set_run_font(run, size=10)
    doc.add_paragraph()


def _svg_to_png(svg_path: Path) -> Path | None:
    try:
        from reportlab.graphics import renderPM
        from svglib.svglib import svg2rlg
    except ImportError:
        return None
    if not svg_path.is_file():
        return None
    drawing = svg2rlg(str(svg_path))
    if drawing is None:
        return None
    tmp = Path(tempfile.mkstemp(suffix=".png")[1])
    renderPM.drawToFile(drawing, str(tmp), fmt="PNG", dpi=144)
    return tmp


def _resolve_logo_path(*, logo_svg: Path, logo_png: Path | None) -> tuple[Path | None, bool]:
    """Devuelve (ruta_imagen, es_temporal_para_borrar)."""
    if logo_png is not None and logo_png.is_file():
        return logo_png, False
    if BUNDLED_LOGO_PNG.is_file():
        return BUNDLED_LOGO_PNG, False
    tmp = _svg_to_png(logo_svg)
    if tmp is not None and tmp.is_file():
        return tmp, True
    return None, False


def _find_sudeaseg_reaseguro_xlsx() -> Path | None:
    base = WORKSPACE / "Reportes" / "Carpeta compartida Actuarial"
    if not base.is_dir():
        return None
    for p in base.glob("*.xlsx"):
        n = p.name.lower()
        if "sudeaseg" in n and "reaseguro" in n and not p.name.startswith("~$"):
            return p
    return None


def _fmt_xlsx_val(v) -> str:
    if v is None:
        return ""
    if isinstance(v, datetime):
        return v.strftime("%Y-%m-%d")
    if isinstance(v, date):
        return v.isoformat()
    s = str(v).strip()
    s = re.sub(r"\s+", " ", s)
    if len(s) > 120:
        return s[:117] + "..."
    return s


def _load_reaseguro_sudeaseg_rows(*, max_rows_per_sheet: int = 60) -> tuple[list[list[str]], str]:
    """
    Lee la relación anual de contratos de reaseguro (SUDEASEG) y devuelve filas para tablas Word
    [Moneda, Tipo, Ramo, Capacidad, Retención, Reaseguradores, Participación, Fecha remisión].
    """
    path = _find_sudeaseg_reaseguro_xlsx()
    if path is None:
        return [], "No se localizó el Excel «f- Relación de Contratos de Reaseguro Año 2025 SUDEASEG» en Reportes/Carpeta compartida Actuarial."
    try:
        import openpyxl
    except ImportError:
        return [], "openpyxl no disponible: no se pudo leer el Excel de relación SUDEASEG."

    wb = openpyxl.load_workbook(path, read_only=True, data_only=True)
    all_rows: list[list[str]] = []
    notes: list[str] = []
    for sname in wb.sheetnames:
        if "2025" not in sname:
            continue
        moneda = "VES" if "Bs" in sname else "USD"
        ws = wb[sname]
        hdr_row = None
        start = None
        for i, row in enumerate(ws.iter_rows(values_only=True), 1):
            if row and str(row[0]).strip() == "Tipo de Contrato":
                hdr_row = i
                start = i + 1
                break
        if start is None:
            continue
        count = 0
        for row in ws.iter_rows(min_row=start, values_only=True):
            if not row or row[0] is None:
                continue
            tipo = _fmt_xlsx_val(row[0])
            if not tipo:
                continue
            ramo = _fmt_xlsx_val(row[4]) if len(row) > 4 else ""
            cap = _fmt_xlsx_val(row[5]) if len(row) > 5 else ""
            ret = _fmt_xlsx_val(row[6]) if len(row) > 6 else ""
            reas = _fmt_xlsx_val(row[11]) if len(row) > 11 else ""
            part = _fmt_xlsx_val(row[12]) if len(row) > 12 else ""
            rem = _fmt_xlsx_val(row[14]) if len(row) > 14 else ""
            all_rows.append([moneda, tipo, ramo, cap, ret, reas, part, rem])
            count += 1
            if count >= max_rows_per_sheet:
                notes.append(f"Hoja «{sname.strip()}»: se listan las primeras {max_rows_per_sheet} filas de contrato.")
                break
        notes.append(f"Hoja «{sname.strip()}»: {count} filas volcadas en esta minuta.")
    wb.close()
    note_txt = "Fuente: " + path.as_posix() + ". " + " ".join(notes)
    return all_rows, note_txt


def _find_resumen_pagos_reaseguro_xlsx() -> Path | None:
    for base in (WORKSPACE / "Reportes", WORKSPACE / "Reportes" / "Mesa Actuarial"):
        if not base.is_dir():
            continue
        for p in base.rglob("Resumen de pagos Reaseguros*.xlsx"):
            if not p.name.startswith("~$"):
                return p
    return None


def _fold_key(s: str) -> str:
    return "".join(
        c for c in unicodedata.normalize("NFD", s.lower()) if unicodedata.category(c) != "Mn"
    )[:52]


def _reas_nombres_desde_celda(cel) -> list[str]:
    if cel is None:
        return []
    out: list[str] = []
    for ln in re.split(r"[\n;]+", str(cel)):
        ln = re.sub(r"\s+", " ", ln).strip()
        if len(ln) < 8:
            continue
        if re.fullmatch(r"[\d%\s.,/-]+", ln):
            continue
        ln = ln.replace("Redbrige", "Redbridge")
        out.append(ln[:58])
    return out[:6] if out else []


def _load_tabla_participacion_estimada_usd() -> tuple[list[list[str]], str]:
    """
    Reparte el total de primas cedidas en USD entre contrapartes según su peso en líneas
    de la relación declarada al supervisor (aproximación para lectura ejecutiva).
    """
    path = _find_sudeaseg_reaseguro_xlsx()
    tot = _try_total_usd_numeric()
    if path is None or tot is None:
        return [], "No se pudo armar la tabla de participación (falta relación en USD o total de primas cedidas)."

    try:
        import openpyxl
    except ImportError:
        return [], "openpyxl no disponible."

    from collections import defaultdict

    acc: dict[str, dict] = defaultdict(lambda: {"w": 0.0, "ramos": set(), "label": ""})
    wb = None
    try:
        wb = openpyxl.load_workbook(path, read_only=True, data_only=True)
        for sname in wb.sheetnames:
            if "2025" not in sname or "Bs" in sname:
                continue
            ws = wb[sname]
            start = None
            for row in ws.iter_rows(values_only=True):
                if row and str(row[0]).strip() == "Tipo de Contrato":
                    start = True
                    continue
                if start is None:
                    continue
                if not row or row[0] is None:
                    continue
                names = _reas_nombres_desde_celda(row[11] if len(row) > 11 else None)
                if not names:
                    names = ["(Varias contrapartes en la línea)"]
                w_each = 1.0 / len(names)
                ramo = str(row[4]).strip()[:45] if len(row) > 4 and row[4] else "—"
                for n in names:
                    key = _fold_key(n.strip()[:52])
                    if not acc[key]["label"]:
                        acc[key]["label"] = n.strip()[:58]
                    acc[key]["w"] += w_each
                    acc[key]["ramos"].add(ramo)
    except OSError:
        return [], "No se pudo leer el Excel de relación."
    finally:
        if wb is not None:
            try:
                wb.close()
            except OSError:
                pass

    tw = sum(d["w"] for d in acc.values()) or 1.0
    rows_out: list[list[str]] = []
    for _, d in sorted(acc.items(), key=lambda kv: -kv[1]["w"])[:12]:
        pct = 100.0 * d["w"] / tw
        monto = tot * (d["w"] / tw)
        rlist = sorted(d["ramos"])
        ramtxt = ", ".join(rlist[:5])
        if len(rlist) > 5:
            ramtxt += "…"
        rows_out.append(
            [
                d["label"],
                f"≈ {monto:,.0f} USD",
                f"≈ {pct:.1f} %",
                ramtxt or "—",
            ]
        )
    nota = (
        "**Lectura orientativa:** se reparte el **total de primas cedidas en USD** del **control interno** entre "
        "contrapartes según su **peso en líneas** de la **relación** al supervisor; **no** sustituye el cierre de **Finanzas**."
    )
    return rows_out, nota


def _try_leer_total_usd_pagos_reaseguro() -> str:
    p = _find_resumen_pagos_reaseguro_xlsx()
    if p is None:
        return "—"
    try:
        import openpyxl

        wb = openpyxl.load_workbook(p, read_only=True, data_only=True)
        ws = wb[wb.sheetnames[0]]
        for row in ws.iter_rows(values_only=True):
            if not row or len(row) < 3:
                continue
            etiqueta = row[1]
            if etiqueta is None:
                continue
            if "TOTAL PAGOS EN USD" in str(etiqueta).upper():
                v = row[2]
                wb.close()
                if isinstance(v, (int, float)):
                    return f"≈ {v:,.0f} USD"
                return str(v)
        wb.close()
    except OSError:
        return "—"
    return "—"


def _try_total_usd_numeric() -> float | None:
    """Devuelve el total de pagos/primas cedidas en USD como número, si se puede leer del Excel de resumen."""
    s = _try_leer_total_usd_pagos_reaseguro()
    if s == "—":
        return None
    digits = re.sub(r"[^\d]", "", s)
    if len(digits) < 4:
        return None
    try:
        return float(digits)
    except ValueError:
        return None


def _distill_reasegurador_cell(
    cells: list[str],
    *,
    max_snippets: int = 10,
    max_line_len: int = 110,
) -> str:
    """Resume celdas de contraparte en una sola lectura legible (sin nombres de archivo)."""
    snippets: list[str] = []
    seen: set[str] = set()
    for raw in cells:
        if not raw:
            continue
        s = str(raw).replace("\r", "")
        for ln in re.split(r"[\n]+", s):
            ln = re.sub(r"\s+", " ", ln).strip()
            if len(ln) < 10:
                continue
            low = ln.lower()
            if low in seen:
                continue
            seen.add(low)
            snippets.append(ln[:max_line_len])
            if len(snippets) >= max_snippets:
                break
        if len(snippets) >= max_snippets:
            break
    return " · ".join(snippets) if snippets else "—"


def _truncar_ejecutivo(s: str, max_chars: int = 380) -> str:
    s = s.strip()
    if len(s) <= max_chars:
        return s
    cut = s[: max_chars - 1]
    sp = cut.rfind(" ")
    return (cut if sp < 20 else cut[:sp]) + "…"


def _load_reaseguro_sintesis_por_moneda() -> tuple[list[list[str]], str]:
    """
    Dos filas como máximo (USD / VES): lectura ejecutiva sin desglose por ramo ni referencias a archivos.
    """
    path = _find_sudeaseg_reaseguro_xlsx()
    if path is None:
        return [], "No se dispuso de la información de reaseguro declarada al supervisor para esta síntesis."

    try:
        import openpyxl
    except ImportError:
        return [], "openpyxl no disponible."

    from collections import defaultdict

    by_moneda: dict[str, list[str]] = defaultdict(list)
    wb = openpyxl.load_workbook(path, read_only=True, data_only=True)
    for sname in wb.sheetnames:
        if "2025" not in sname:
            continue
        moneda = "Dólares" if "Bs" not in sname else "Bolívares"
        ws = wb[sname]
        start = None
        for i, row in enumerate(ws.iter_rows(values_only=True), 1):
            if row and str(row[0]).strip() == "Tipo de Contrato":
                start = i + 1
                break
        if start is None:
            continue
        for row in ws.iter_rows(min_row=start, values_only=True):
            if not row or row[0] is None:
                continue
            cel = row[11] if len(row) > 11 else None
            if cel:
                by_moneda[moneda].append(str(cel))
    wb.close()

    rows: list[list[str]] = []
    for mon, label_hint in (
        ("Dólares", "Salud de alto costo, vida y daños con **panel internacional**; fianzas y automóvil con contrapartes recurrentes."),
        ("Bolívares", "Daños generales, fianzas y ramos locales con **contrapartes mayoritariamente domésticas**."),
    ):
        if mon not in by_moneda:
            continue
        nombres = _distill_reasegurador_cell(by_moneda[mon], max_snippets=4, max_line_len=72)
        if nombres == "—":
            fila = [mon, _truncar_ejecutivo(label_hint, 320)]
        else:
            fila = [
                mon,
                _truncar_ejecutivo(f"{label_hint} **Contrapartes que figuran** en la información al supervisor: {nombres}.", 420),
            ]
        rows.append(fila)

    tot = _try_leer_total_usd_pagos_reaseguro()
    nota = (
        "Síntesis **cualitativa** por moneda; el **detalle contractual** y el **cierre de cifras** corresponden a **Legal** y **Finanzas**. "
        f"Referencia interna de magnitud (**primas cedidas en USD**): **{tot}**."
    )
    return rows, nota


def _redbridge_participacion_resumen() -> str:
    """
    Resume participaciones declaradas donde aparece Redbridge (hojas USD del Excel de relación 2025).
    Devuelve texto corto para incrustar en párrafo; cadena vacía si no hay datos.
    """
    path = _find_sudeaseg_reaseguro_xlsx()
    if path is None:
        return ""
    try:
        import openpyxl
    except ImportError:
        return ""

    pcts: list[int] = []
    cede_100 = False
    wb = None
    try:
        wb = openpyxl.load_workbook(path, read_only=True, data_only=True)
        for sname in wb.sheetnames:
            if "2025" not in sname or "Bs" in sname:
                continue
            ws = wb[sname]
            start = None
            for row in ws.iter_rows(values_only=True):
                if row and str(row[0]).strip() == "Tipo de Contrato":
                    start = True
                    continue
                if start is None:
                    continue
                if not row or row[0] is None:
                    continue
                blob = " ".join(str(x or "") for x in row[:15])
                if "redbridge" not in blob.lower():
                    continue
                for m in re.finditer(r"(\d+)\s*%", blob):
                    pcts.append(int(m.group(1)))
                if re.search(r"\b100\s*%", blob):
                    cede_100 = True
    except OSError:
        return ""
    finally:
        if wb is not None:
            try:
                wb.close()
            except OSError:
                pass

    partes: list[str] = []
    if pcts:
        partes.append(
            f"en las líneas donde figura Redbridge aparecen participaciones de hasta {max(pcts)}% "
            f"(máximo observado en la relación revisada)"
        )
    if cede_100:
        partes.append(
            "constan además tramos en USD con 100% de participación declarada, lo que implica cedencia total "
            "en esos contratos y alto acople si la contraparte falla"
        )
    if not partes:
        return ""
    return "; ".join(partes) + "."


def _parrafo_reaseguro_sintesis_unica() -> str:
    """Un solo texto ejecutivo: esquemas, ramos sensibles, contrapartes y magnitud (sin tablas)."""
    tot = _try_leer_total_usd_pagos_reaseguro()
    rb_line = _redbridge_participacion_resumen()
    path = _find_sudeaseg_reaseguro_xlsx()
    if path is None:
        return _truncar_ejecutivo(
            "La compañía declara **reaseguro** en **bolívares y dólares** con esquemas de **cuota parte**, **excedente** "
            "y tramos **facultativos** o **no proporcionales** en **salud de alto costo**, **vida** y ramos patrimoniales; "
            "**Redbridge** concentra **riesgo de contraparte** en **divisas**. "
            f"Referencia de magnitud (**primas cedidas en USD**): **{tot}**.",
            620,
        )
    try:
        import openpyxl
    except ImportError:
        return (
            "Declaración **anual** de reaseguro ante el supervisor en **dos monedas**, con foco en **salud** y **vida** en **USD**. "
            f"Referencia de magnitud: **{tot}**."
        )

    all_cels: list[str] = []
    wb = openpyxl.load_workbook(path, read_only=True, data_only=True)
    for sname in wb.sheetnames:
        if "2025" not in sname:
            continue
        ws = wb[sname]
        start = None
        for i, row in enumerate(ws.iter_rows(values_only=True), 1):
            if row and str(row[0]).strip() == "Tipo de Contrato":
                start = i + 1
                break
        if start is None:
            continue
        for row in ws.iter_rows(min_row=start, values_only=True):
            if not row or row[0] is None:
                continue
            cel = row[11] if len(row) > 11 else None
            if cel:
                all_cels.append(str(cel))
    wb.close()

    nombres = _distill_reasegurador_cell(all_cels, max_snippets=5, max_line_len=48)
    core = (
        "La **relación anual** al supervisor resume **cuota parte** y **excedente** en patrimoniales y **RC**, "
        "más **facultativo** y **capas** en **salud de alto costo (HCM)**, **vida**, automóvil y fianzas; **Redbridge** "
        "aparece de forma **dominante** en **líneas en divisas**. "
    )
    if rb_line:
        core += f"**Concentración:** {rb_line} "
    if nombres != "—":
        core += f"Otras contrapartes citadas en la misma remisión: {nombres}. "
    core += f"Referencia de magnitud (**primas cedidas en USD**): **{tot}**."
    return _truncar_ejecutivo(core, 720)


def _append_bloque_informe_confidencial_conciso(doc: Document) -> None:
    """Versión breve del informe confidencial: sin subapartados 3.1–3.3 ni tablas regulatorias detalladas."""
    _heading(doc, "1. Hallazgos y contexto", level=1)
    _p_rich(
        doc,
        "Desde el **trabajo actuarial** se revisó, a **nivel general**, la **documentación legal** de apoyo, "
        "lo que permitió **focalizar** el análisis en **reaseguro** y **cumplimiento** ante el supervisor. "
        "Destaca la **Comunicación SAA-04-1443** (**febrero 2026**) por **remisión extemporánea** de **información contractual** "
        "de reaseguro. La **respuesta** al regulador debe ir **coordinada** entre **Actuarial**, el **área de reaseguro** y **Legal**.",
        size=11,
    )
    _heading(doc, "2. Reaseguro y supervisor", level=1)
    _p_rich(
        doc,
        "En **dólares**, la **información al supervisor** se resume mejor con **quién pesa** y **qué parte** del "
        "**total de primas cedidas** (cifra de **control interno**) correspondería a cada contraparte, según el "
        "**peso en líneas** del archivo regulador (**aproximación** para **junta**, no cierre contable).",
        size=10,
    )
    part_rows, part_note = _load_tabla_participacion_estimada_usd()
    _p_rich(doc, part_note, size=9)
    if part_rows:
        _add_table(
            doc,
            ["Contraparte (USD)", "Monto orientativo (USD)", "% del total informado", "Ramos (declarados)"],
            part_rows,
        )
    rb = _redbridge_participacion_resumen()
    if rb:
        _p_rich(doc, f"**Detalle Redbridge:** {rb.strip().rstrip('.')}.", size=10)
    _p_rich(
        doc,
        "**PAS** (**Procedimiento Administrativo Sancionatorio**): el retraso formal incrementa el riesgo de **sanción** "
        "y de **mayor escrutinio**; la **defensa técnica** y la **consistencia** de la información deben alinearse entre "
        "**Actuarial**, **reaseguro** y **Legal**.",
        size=11,
    )
    _heading(doc, "3. Riesgos en foco", level=1)
    filas = _tabla_hallazgos_reaseguro_anexo_ejecutivo()
    for row in filas[1:]:
        if len(row) < 3:
            continue
        riesgo, tema, lectura = row[0], row[1], row[2]
        p = doc.add_paragraph()
        p.paragraph_format.space_after = Pt(6)
        r0 = p.add_run(f"{riesgo} — {tema}. ")
        _set_run_font(r0, bold=True, size=10)
        r1 = p.add_run(re.sub(r"\*\*", "", lectura))
        _set_run_font(r1, size=10)


def _load_reaseguro_contrapartes_por_ramos_rows() -> tuple[list[list[str]], str]:
    """
    Vista muy resumida: ramos agrupados + contrapartes que aparecen en la **información enviada**
    al supervisor (sin tabla operativa línea a línea).
    """
    path = _find_sudeaseg_reaseguro_xlsx()
    if path is None:
        return [], "No se pudo leer la información contractual de reaseguro enviada al regulador (archivo no localizado)."

    try:
        import openpyxl
    except ImportError:
        return [], "openpyxl no disponible."

    from collections import defaultdict

    bucket_cells: dict[str, list[str]] = defaultdict(list)
    wb = openpyxl.load_workbook(path, read_only=True, data_only=True)
    for sname in wb.sheetnames:
        if "2025" not in sname:
            continue
        moneda = "Bolívares" if "Bs" in sname else "Dólares"
        ws = wb[sname]
        start = None
        for i, row in enumerate(ws.iter_rows(values_only=True), 1):
            if row and str(row[0]).strip() == "Tipo de Contrato":
                start = i + 1
                break
        if start is None:
            continue
        for row in ws.iter_rows(min_row=start, values_only=True):
            if not row or row[0] is None:
                continue
            tipo = str(row[0]).strip()
            ramo = str(row[4]).strip() if len(row) > 4 and row[4] else ""
            rl = ramo.lower()
            if moneda == "Dólares" and (
                "hcm" in rl
                or "zafiro" in rl
                or "diamante" in rl
                or "salud" in rl
                or "no prop" in tipo.lower()
                or "no proporcional" in tipo.lower()
            ):
                fam = "Salud (capas y productos de alto costo)"
            elif moneda == "Dólares" and "vida" in rl:
                fam = "Vida (colectivo e individual)"
            elif moneda == "Dólares":
                fam = "Daños, RC, transporte y ramos técnicos afines"
            elif "fianza" in rl or "auto" in rl or "casco" in rl or "fac" in tipo.lower():
                fam = "Fianzas, automóvil y facultativos obligatorios"
            else:
                fam = "Daños y ramos generales"
            key = f"{moneda} — {fam}"
            cel = row[11] if len(row) > 11 else None
            if cel:
                bucket_cells[key].append(str(cel))
    wb.close()

    rows: list[list[str]] = []
    for key in sorted(bucket_cells.keys(), key=lambda k: (not k.startswith("Dólares"), k)):
        rows.append([key, _distill_reasegurador_cell(bucket_cells[key])])

    tot_usd = _try_leer_total_usd_pagos_reaseguro()
    nota = (
        "La tabla resume **quién aparece** en la **información de reaseguro enviada** al supervisor, "
        f"agrupando **ramos**. Como referencia de magnitud, el **control interno de tesorería** sitúa las "
        f"primas cedidas en dólares en torno a **{tot_usd}** (cifra orientativa; **Finanzas** confirma el cierre)."
    )
    return rows, nota


def _load_reaseguro_resumen_tecnico_rows() -> tuple[list[list[str]], str]:
    """
    Resumen por moneda y familia de contrato (cuatro columnas) para uso interno
    (``para_accionista=False``).
    """
    path = _find_sudeaseg_reaseguro_xlsx()
    if path is None:
        return [], "No se localizó el Excel de relación SUDEASEG 2025."

    try:
        import openpyxl
    except ImportError:
        return [], "openpyxl no disponible."

    from collections import defaultdict

    grupos: dict[tuple[str, str], int] = defaultdict(int)
    wb = openpyxl.load_workbook(path, read_only=True, data_only=True)
    for sname in wb.sheetnames:
        if "2025" not in sname:
            continue
        moneda = "Bolívares" if "Bs" in sname else "Dólares"
        ws = wb[sname]
        start = None
        for i, row in enumerate(ws.iter_rows(values_only=True), 1):
            if row and str(row[0]).strip() == "Tipo de Contrato":
                start = i + 1
                break
        if start is None:
            continue
        for row in ws.iter_rows(min_row=start, values_only=True):
            if not row or row[0] is None:
                continue
            tipo = str(row[0]).strip()
            ramo = str(row[4]).strip() if len(row) > 4 and row[4] else ""
            rl = ramo.lower()
            if moneda == "Dólares" and (
                "hcm" in rl
                or "zafiro" in rl
                or "diamante" in rl
                or "salud" in rl
                or "no prop" in tipo.lower()
                or "no proporcional" in tipo.lower()
            ):
                fam = "Salud (capas no proporcionales / HCM / productos internacionales)"
            elif moneda == "Dólares" and "vida" in rl:
                fam = "Vida (colectivo e individual)"
            elif moneda == "Dólares":
                fam = "Daños, responsabilidad civil, transporte y ramos técnicos afines"
            elif "fianza" in rl or "auto" in rl or "casco" in rl or "fac" in tipo.lower():
                fam = "Fianzas, automóvil y facultativos obligatorios (bolívares)"
            else:
                fam = "Daños y ramos generales (bolívares)"
            grupos[(moneda, fam)] += 1
    wb.close()

    rows: list[list[str]] = []
    for (mon, fam) in sorted(grupos.keys(), key=lambda x: (x[0] != "Dólares", x[1])):
        n = grupos[(mon, fam)]
        if mon == "Dólares" and "Salud" in fam:
            riesgo = (
                "Aquí suele concentrarse mucho del gasto en divisas; la contraparte Redbridge "
                "aparece de forma recurrente en pagos y en contratos: alto acople si ese grupo tuviera tensiones."
            )
        elif mon == "Dólares":
            riesgo = (
                "Riesgo repartido entre varios reaseguradores (p. ej. Bee, Delta, Kairos, R.I.V.); "
                "dependencia moderada frente a la línea de salud."
            )
        else:
            riesgo = "Principalmente contrapartes locales; exposición en bolívares acotada frente al total en USD."
        rows.append([mon, fam, str(n), riesgo])

    tot_usd = _try_leer_total_usd_pagos_reaseguro()
    nota = (
        f"Fuente relación contractual: {path.as_posix()}. "
        f"Referencia de magnitud de pagos 2025 (prima neta cedida, archivo interno): total declarado en resumen **{tot_usd}** "
        "(validar con Finanzas)."
    )
    return rows, nota


def _tabla_hallazgos_semana1_y2() -> list[list[str]]:
    """Riesgo primero (formato lámina Mesa Actuarial Semana 1 PDF), complementada con Semana 2."""
    return [
        ["Riesgo", "Ámbito", "Hallazgo — Semana 1 (diagnóstico inicial) / complemento Semana 2"],
        [
            "Medio / Bajo",
            "Reaseguro — cuerpo contractual",
            "S1: material resumido de flujos; faltaba cuerpo contractual integrado. "
            "S2: PDF replicados en Mesa (`vaciado-completo-2026-04-23` y `01-Reaseguro-Contratos`); "
            "matriz `MATRIZ-REASEGURO-Legal-vs-SUDEASEG-2026-04-24.md`; cruce con relación SUDEASEG 2025.",
        ],
        [
            "Medio / Alto",
            "Reportería y BI",
            "S1: reportería manual con archivos Excel; urgencia de BI. "
            "S2: consolidación manual aceptable en la semana; **sigue sin** tubería automatizada a Mesa ni tableros BI.",
        ],
        [
            "Medio",
            "Productos / SUDEASEG",
            "S1: cola de productos pendiente de aprobación. "
            "S2: en espera de confirmación de entrega; condicionados replicados en `De-Gerencia-Legal/02-Productos-Condicionados-SUDEASEG`.",
        ],
        [
            "Medio",
            "Solvencia / MS",
            "S1: paquete de solvencia incompleto en oleada inicial. "
            "S2: sin cambio sustantivo declarado en esta minuta; seguir pliego MS-01/02 con Finanzas.",
        ],
        [
            "Medio / Bajo",
            "Certificación de reservas",
            "S1: pendiente de incorporación. "
            "S2: **cubierto** en carpeta compartida (PDF certificación, ONR, reserva complementaria).",
        ],
        [
            "Medio",
            "Carga analítica al regulador",
            "S1: pendiente evidencia de cuadre. "
            "S2: recibos de prima cubiertos; **pendiente** inventario de demás formatos analíticos con Tecnología/Finanzas.",
        ],
        [
            "Alto",
            "Procedimiento administrativo — reaseguro",
            "S1: no explicitado en lámina. "
            "S2: consta **Comunicación SAA-04-1443** (27-feb-2026) por **remisión extemporánea de contratos de reaseguro**; "
            "la documentación correspondiente obra en el **expediente de procedimientos** atendido por **Legal**. "
            "Ver § reaseguro y PAS en esta minuta.",
        ],
        [
            "Medio",
            "Lectura sectorial vs libros",
            "S1: convergencia conceptual; falta puente metodológico. "
            "S2: no fue foco operativo de la semana 2; sigue como tema de profundización en oleadas posteriores.",
        ],
        [
            "Bajo",
            "Integración documental (Semana 2)",
            "S2 únicamente: promoción **176** archivos actuarial + espejo Legal **176** archivos; índices de promoción y vaciado; "
            "minuta y presentación de control ampliadas (7 láminas).",
        ],
    ]


def _tabla_hallazgos_mesa_sin_detalle_reaseguro() -> list[list[str]]:
    """Misma estructura que la lámina Mesa Actuarial, sin filas de detalle contractual ni PAS (van al informe CONF)."""
    return [
        ["Riesgo", "Ámbito", "Hallazgo — Semana 1 / complemento Semana 2"],
        [
            "Medio / Alto",
            "Reportería y BI",
            "S1: reportería manual; urgencia de BI. S2: sin tubería automatizada a Mesa ni tableros BI.",
        ],
        [
            "Medio",
            "Productos / SUDEASEG",
            "S2: cola en espera; condicionados en `De-Gerencia-Legal/02-Productos-Condicionados-SUDEASEG`.",
        ],
        [
            "Medio",
            "Solvencia / MS",
            "S1: paquete incompleto. S2: seguir MS-01/02 con Finanzas.",
        ],
        [
            "Medio / Bajo",
            "Certificación de reservas",
            "S2: cubierto en carpeta compartida (certificación, ONR, reserva complementaria).",
        ],
        [
            "Medio",
            "Carga analítica al regulador",
            "S2: recibos cubiertos; pendiente inventario de otros formatos con Tecnología/Finanzas.",
        ],
        [
            "Medio",
            "Lectura sectorial vs libros",
            "S1: falta puente metodológico. S2: profundización en oleadas posteriores.",
        ],
        [
            "Bajo",
            "Integración documental (Semana 2)",
            "S2: promoción 176 archivos actuarial + espejo Legal 176; índices; minuta y PPTX 7 láminas.",
        ],
        [
            "—",
            "Reaseguro / TPA salud / PAS",
            "**Detalle retirado de esta minuta general.** Ver informe confidencial "
            "`CONF-Accionista-Reaseguro-TPA-Salud-2026-04-24.docx` (distribución restringida).",
        ],
    ]


def _tabla_hallazgos_reaseguro_anexo_confidencial() -> list[list[str]]:
    """Subconjunto para el anexo CONF: reaseguro, PAS e integración documental (Semana 2)."""
    full = _tabla_hallazgos_semana1_y2()
    idxs = [0, 1, 7, 9]  # encabezado + reaseguro + PAS + integración
    return [full[i] for i in idxs if i < len(full)]


def _tabla_hallazgos_reaseguro_anexo_ejecutivo() -> list[list[str]]:
    """«Riesgo primero», dos filas: contrato faltante y PAS (sin fila operativa de cruce legal)."""
    return [
        ["Riesgo", "Tema", "Lectura para junta"],
        [
            "Medio / Bajo",
            "Contratos de reaseguro",
            "Persiste **laguna de respaldo** respecto del **contrato no proporcional** vinculado a un tramo del orden de **~4,3 MUSD** "
            "(orden de magnitud sujeto a **confirmación**); conviene **cerrar expediente** con **Legal** y **reaseguro** "
            "antes de ampliar exposición.",
        ],
        [
            "Alto",
            "Procedimiento del supervisor (PAS)",
            "**PAS** (**Procedimiento Administrativo Sancionatorio**): trámite abierto por **remisión extemporánea** de la "
            "**información contractual de reaseguro** (**SAA-04-1443**); la **contención** del riesgo reputacional y sancionatorio "
            "requiere **respuesta coordinada** entre **Actuarial**, **reaseguro** y **Legal**.",
        ],
    ]


def append_confidential_reinsurance_tpa_sections(
    doc: Document,
    *,
    para_accionista: bool = True,
    incluir_matriz_detallada: bool = False,
    modo_informe_ejecutivo: bool = False,
) -> None:
    """
    Bloque retirado de la minuta general. Si ``para_accionista`` es True, tono ejecutivo,
    negritas coherentes (** vía ``_p_rich`` / tablas) y **contrapartes / ramos agrupados**
    a partir de la información enviada al supervisor (salvo ``incluir_matriz_detallada``).

    Si ``modo_informe_ejecutivo`` es True (solo informe confidencial Word), se sustituye todo el bloque
    por una **versión breve** sin tablas detalladas de esquemas ni subapartados 3.1–3.3.
    """
    if modo_informe_ejecutivo and para_accionista:
        _append_bloque_informe_confidencial_conciso(doc)
        return

    _heading(doc, "1. Confidencialidad y uso interno", level=1)
    if para_accionista:
        _p_rich(
            doc,
            "Este documento reúne **información sensible**: cómo la compañía **transfiere riesgos** a otras "
            "aseguradoras (reaseguro), cómo se administra la **red de salud** con terceros (modelo tipo **TPA**) "
            "y un **procedimiento del supervisor** por **envío fuera de plazo** de contratos de reaseguro. "
            "Solo debe circular entre **junta directiva**, **accionariado de referencia** autorizado, **Legal** "
            "y **Actuarial**; no es material para reuniones abiertas ni para el paquete público del POR.",
        )
    else:
        _p_rich(
            doc,
            "Este documento contiene **información sensible** sobre estructura de reaseguro, contrapartes, "
            "montos y procedimiento administrativo sancionatorio. Su circulación debe limitarse a **alta dirección**, "
            "**asesoría externa autorizada** y **Gerencia Legal / Actuarial** según política interna de La Internacional. "
            "No forma parte del paquete de divulgación general del POR.",
        )

    _heading(doc, "2. Antecedentes y revisión legal", level=1)
    if para_accionista:
        _p_rich(
            doc,
            "Ante **indicios de anomalía** en la **información de soporte** del **ciclo actuarial**, se procedió a una "
            "**revisión coordinada del área legal** para **completar el análisis**. De ese cruce surgieron **hallazgos "
            "relevantes y sensibles** que este informe sintetiza, con foco en **reaseguro** y **supervisor**.",
        )
    else:
        _p(
            doc,
            "Con fecha de espejo **2026-04-23**, se replicó el contenido de `Reportes/GERENCIA LEGAL` en "
            "`Reportes/Mesa Actuarial/Gerencia-Legal/vaciado-completo-2026-04-23/` **sin modificar el origen**. "
            "Para el flujo de trabajo actuarial se clasificaron copias bajo "
            "`Reportes/Mesa Actuarial/Materiales/De-Gerencia-Legal/`, con el mapa "
            "`Reportes/Mesa Actuarial/Gerencia-Legal/INDICE-VACIADO-Y-MAPA-A-MESA-ACTUARIAL.md`. "
            "Estructura temática (resumen):",
        )
    if para_accionista:
        _p_rich(
            doc,
            "El cruce con **Legal** priorizó **reaseguro y supervisor**, donde destaca la **Comunicación SAA-04-1443** "
            "por **información contractual fuera de plazo**, y dejó acotados los frentes de **oferta y canal**, "
            "**cumplimiento** y **contratación de soporte** que inciden en el **perfil de riesgo** y la **reputación** "
            "ante el ente regulador.",
        )
    else:
        for text in [
            "**Contratos de reaseguro** (PDF) y su cruce con lo declarado al regulador.",
            "**Productos y condicionados** frente a SUDEASEG (oferta, oficios).",
            "**Intermediarios**, canales alternativos e incentivos comerciales.",
            "**Procedimientos sancionatorios** — aquí está la comunicación **SAA-04-1443** sobre el retraso en informar contratos de reaseguro.",
            "**Cumplimiento** (prevención de lavado de activos), **juicios**, **contribuciones** y **solvencias**.",
            "**Contratos generales** (arrendamiento, proveedores, ISR, propiedad intelectual).",
        ]:
            _bullet_rich(doc, text)

    _heading(doc, "3. Reaseguro y supervisor (declaraciones y cumplimiento)", level=1)

    _heading(doc, "3.1. Qué declara la compañía al regulador (2025)", level=2)
    if para_accionista:
        _p_rich(
            doc,
            "La **Superintendencia** exige una **relación anual** de los **contratos de reaseguro**: con **quién**, "
            "hasta **qué montos** y en **qué moneda** se cede riesgo. La **información enviada** para **2025** describe, "
            "en términos generales, **tratados de reparto** (excedente y cuota parte) en **dólares y bolívares** para "
            "daños y ramos afines; **capas** para **salud** (productos de alto costo y esquemas con terceros administradores); "
            "**vida** con contrapartes internacionales; y **fianzas y automóvil** con presencia recurrente de **Redbridge** "
            "en varias líneas. Los **contratos firmados** forman parte del **respaldo** ya incorporado al análisis.",
        )
    else:
        _p(
            doc,
            "Según la relación anual remitida al regulador (archivo "
            "«f- Relación de Contratos de Reaseguro Año 2025 SUDEASEG.xlsx» en `Reportes/Carpeta compartida Actuarial`), "
            "la compañía declara, entre otros, **tratados en excedente de pérdida** y **cuota parte** en bolívares "
            "y en dólares; **facultativo obligatorio** en ramos como fianzas y automóvil casco (contraparte "
            "Redbridge en varias líneas en USD); **vida** con panel Kairos / Hannover y corretaje AON; y capas "
            "**no proporcionales** en salud (HCM / productos Zafiro y Diamante) con fechas de remisión o notas "
            "de endoso en el propio Excel. Los **PDF** de contrato vigentes para trabajo de mesa están en "
            "`Materiales/De-Gerencia-Legal/01-Reaseguro-Contratos/` (programa Bouquet 2026, proporcionales auto/fianzas/HCM, "
            "vida AON).",
        )

    _heading(
        doc,
        "3.2. Contrapartes en la información al supervisor (síntesis por moneda)"
        if para_accionista
        else "3.2. Matriz de reaseguro — vista concentrada (sustituye la tabla línea por línea)",
        level=2,
    )
    if para_accionista:
        res_rows, res_note = _load_reaseguro_sintesis_por_moneda()
        _p_rich(doc, res_note, size=10)
    else:
        res_rows, res_note = _load_reaseguro_resumen_tecnico_rows()
        _p_rich(doc, res_note, size=10)
    if res_rows:
        hdr = (
            ["Moneda", "Síntesis"]
            if para_accionista
            else [
                "Moneda",
                "Bloque de negocio (agrupado)",
                "Líneas en la relación 2025",
                "Qué implica para el riesgo de la compañía",
            ]
        )
        _add_table(doc, hdr, res_rows)
    else:
        _p(doc, "No se pudo generar el resumen automático.", size=10)
    if para_accionista:
        _p_rich(
            doc,
            "El **cuadre** con **contratos** y la **respuesta** al supervisor siguen bajo responsabilidad de **Legal** "
            "con el **apoyo técnico** que corresponda.",
            size=10,
        )
    else:
        _p_rich(
            doc,
            "El **detalle fila a fila** (cada ramo y cada contraparte) sigue disponible en el **mismo Excel** "
            "y en la matriz técnica **Legal vs SUDEASEG** citada más abajo; aquí solo se **condensa** para lectura de junta.",
            size=10,
        )

    if incluir_matriz_detallada:
        if para_accionista:
            _p_rich(
                doc,
                "En esta **versión ejecutiva** no se incorpora el **detalle línea a línea** del reaseguro; queda "
                "disponible para **uso interno** si **Legal** o **Actuarial** lo requieren.",
                size=9,
            )
        else:
            _heading(doc, "3.2 bis. Detalle regulatorio línea a línea (opcional)", level=2)
            reas_rows, reas_note = _load_reaseguro_sudeaseg_rows()
            _p(doc, reas_note, italic=True, size=9)
            if reas_rows:
                _add_table(
                    doc,
                    [
                        "Moneda",
                        "Tipo de contrato",
                        "Ramo",
                        "Capacidad",
                        "Retención",
                        "Reasegurador(es)",
                        "Participación",
                        "Fecha remisión SUDEASEG",
                    ],
                    reas_rows,
                )

    _heading(doc, "3.3. Por qué es delicado no presentar el contrato de reaseguro a tiempo (SAA-04-1443)", level=2)
    _p_rich(
        doc,
        "El supervisor necesita ver **a tiempo** con **quién** y **bajo qué reglas** la compañía cede riesgo. "
        "Si la información llega **tarde**, no es un simple “tramite”: abre un **procedimiento administrativo "
        "sancionatorio** (PAS). En la práctica eso significa **multas** (la ley nueva del sector habla de rangos "
        "muy altos expresados en función del **tipo de cambio de referencia**), **presión reputacional** frente al "
        "**mercado y los accionistas**, y **mayor escrutinio** en futuras revisiones. En este expediente consta la "
        "**Comunicación SAA-04-1443** (febrero 2026) por **remisión extemporánea** de **contratos de reaseguro**; "
        "la **documentación** asociada consta en el **expediente de procedimientos** que coordina **Legal**.",
    )
    _p_rich(
        doc,
        "**Lectura de riesgo para la compañía:** si la compañía depende de unas pocas contrapartes en **salud en dólares** "
        "y además aparece un **retraso formal** ante el regulador, se combinan **riesgo de contraparte** (¿qué pasa "
        "si ese grupo financiero falla?) y **riesgo de cumplimiento** (multas y desconfianza del supervisor). "
        "No implica pronunciamiento jurídico: el **oficio completo** debe revisarlo **Legal** con el **cuerpo del contrato** "
        "y la **documentación de respaldo** que corresponda con **Finanzas** y **Actuarial**.",
    )

    if para_accionista:
        _heading(doc, "4. Material de apoyo interno", level=1)
        for text in [
            "Cruce entre **contratos firmados** y lo **declarado al supervisor**.",
            "Relación de **montos cedidos** y calendario **administrativo** (PAS).",
            "Presentación breve para **conversaciones reservadas** de junta.",
        ]:
            _bullet_rich(doc, text, size=10)
    else:
        _heading(doc, "4. Referencias técnicas en `Mesa Actuarial/Materiales`", level=1)
        for text in [
            "**MATRIZ-REASEGURO-Legal-vs-SUDEASEG-2026-04-24.md** — cruce contratos PDF vs relación SUDEASEG.",
            "**MATRIZ-REASEGURO-EJECUTIVA-MONTOS-CONTRATOS-Y-PA-2026-04-23.md** — primas pagadas y marco PAS.",
            "**La_Internacional_Reaseguro_Vista_Accionista_2026.pptx** — láminas de apoyo para conversación restringida.",
        ]:
            _bullet_rich(doc, text, size=10)

    if para_accionista:
        _heading(doc, "5. Síntesis de riesgos (reaseguro y PAS)", level=1)
        _p_rich(
            doc,
            "Se mantiene el criterio corporativo **riesgo primero**, **acotado** aquí a **reaseguro** y **PAS**.",
            size=10,
        )
        filas = _tabla_hallazgos_reaseguro_anexo_ejecutivo()
    else:
        _heading(doc, "5. Tabla integrada (recorte reaseguro / PAS — mismo criterio que Semana 1)", level=1)
        _p_rich(
            doc,
            "Se mantiene el **orden “riesgo primero”** usado en la presentación **La Internacional Semana 1 2026**, "
            "pero **solo** en las líneas que afectan **reaseguro** y **procedimiento**.",
            size=10,
        )
        filas = _tabla_hallazgos_reaseguro_anexo_confidencial()
    _add_table(doc, filas[0], filas[1:])


def _avance_mesa_actuarial_pct(validation_rows: list[list[str]]) -> int:
    """
    Estima el avance global de entrega (Mesa Actuarial) a partir de la columna «Estado».
    Pesos orientativos: Cubierto 100 %, Parcial 55 %, Pendiente / en espera 20 %, Bajo (riesgo) 85 %,
    Alto (riesgo / sin desarrollo) 18 % — promedio simple entre bloques.
    """
    scores: list[int] = []
    for row in validation_rows:
        if len(row) < 2:
            continue
        estado = row[1].strip()
        if estado.startswith("Cubierto"):
            scores.append(100)
        elif estado.startswith("Parcial"):
            scores.append(55)
        elif estado.startswith("Bajo"):
            scores.append(85)
        elif estado.startswith("Alto"):
            scores.append(18)
        elif "Pendiente" in estado or "en espera" in estado.lower():
            scores.append(20)
        else:
            scores.append(45)
    if not scores:
        return 0
    return int(round(sum(scores) / len(scores)))


def _logo_png_bytesio(pic_path: Path) -> io.BytesIO | None:
    """PNG RGB embebible en Word (evita problemas de transparencia / lectura de ruta)."""
    try:
        from PIL import Image
    except ImportError:
        return None
    try:
        im = Image.open(pic_path)
        if im.mode not in ("RGB", "L"):
            im = im.convert("RGB")
        buf = io.BytesIO()
        im.save(buf, format="PNG", optimize=True)
        buf.seek(0)
        return buf
    except OSError:
        return None


def build_document(
    *,
    out_path: Path,
    logo_svg: Path,
    logo_png: Path | None,
    fecha_documento: str,
    elaborado_por: str,
    referencia_por: str,
) -> None:
    doc = Document()
    style = doc.styles["Normal"]
    style.font.name = "Arial"
    style.font.size = Pt(11)
    style._element.rPr.rFonts.set(qn("w:eastAsia"), "Arial")

    pic_path, pic_is_temp = _resolve_logo_path(logo_svg=logo_svg, logo_png=logo_png)
    logo_stream = _logo_png_bytesio(pic_path) if pic_path is not None else None

    # Encabezado: tabla más estrecha que el ancho útil, centrada (evita salida por el margen derecho).
    enc = doc.add_table(rows=1, cols=2)
    enc.autofit = False
    enc.alignment = WD_TABLE_ALIGNMENT.CENTER
    w_logo_col = Inches(3.05)
    w_txt_col = Inches(3.15)
    enc.columns[0].width = w_logo_col
    enc.columns[1].width = w_txt_col
    c_logo = enc.rows[0].cells[0]
    c_txt = enc.rows[0].cells[1]
    c_logo.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
    c_txt.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
    c_logo.width = w_logo_col
    c_txt.width = w_txt_col

    p_logo = c_logo.paragraphs[0]
    p_logo.alignment = WD_ALIGN_PARAGRAPH.LEFT
    p_logo.paragraph_format.space_before = Pt(0)
    p_logo.paragraph_format.space_after = Pt(0)
    if logo_stream is not None:
        p_logo.add_run().add_picture(logo_stream, width=Inches(2.75))
    elif pic_path is not None and pic_path.is_file():
        p_logo.add_run().add_picture(str(pic_path), width=Inches(2.75))
    else:
        r = p_logo.add_run("La Internacional\nde Seguros")
        _set_run_font(r, bold=True, size=8, color=TITLE_BLUE)

    if pic_is_temp and pic_path is not None:
        try:
            pic_path.unlink(missing_ok=True)
        except OSError:
            pass

    pr = c_txt.paragraphs[0]
    pr.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    pr.paragraph_format.space_before = Pt(0)
    pr.paragraph_format.space_after = Pt(0)
    pr.paragraph_format.right_indent = Pt(8)
    r1 = pr.add_run("La Internacional de Seguros, C.A.\n")
    _set_run_font(r1, bold=True, size=9, color=TITLE_BLUE)
    r2 = pr.add_run("Plan de Optimización y Rentabilidad (POR) 2026\n")
    _set_run_font(r2, size=8)
    r3 = pr.add_run("Carpeta compartida Actuarial · Mesa Actuarial")
    _set_run_font(r3, italic=True, size=8, color=ACCENT_GOLD)

    doc.add_paragraph()
    sep = doc.add_paragraph()
    run_sep = sep.add_run("―" * 62)
    _set_run_font(run_sep, size=8, color=ACCENT_GOLD)

    t = doc.add_paragraph()
    t.alignment = WD_ALIGN_PARAGRAPH.CENTER
    tr = t.add_run("MINUTA — MESA ACTUARIAL\n")
    _set_run_font(tr, bold=True, size=18, color=TITLE_BLUE)
    tr2 = t.add_run("Carpeta compartida Actuarial · Validación vs Solicitud consolidada v2")
    _set_run_font(tr2, bold=True, size=12, color=TITLE_BLUE)

    doc.add_paragraph()
    _add_table(
        doc,
        ["Campo", "Detalle"],
        [
            ["Referencia", referencia_por],
            ["Fecha del documento", fecha_documento],
            ["Norma de referencia", "Solicitud consolidada de información «Definitivo v2» (08-abr-2026)"],
            [
                "Ámbito",
                "Carpeta compartida Actuarial; carpeta Mesa Actuarial; vaciado Gerencia Legal bajo "
                "`Reportes/Mesa Actuarial` (sin modificar origen en `Reportes/GERENCIA LEGAL`).",
            ],
            ["Elaborado por", elaborado_por],
        ],
    )

    _heading(doc, "1. Objeto", level=1)
    _p(
        doc,
        "Sintetizar para la alta dirección el contenido aportado en la **carpeta compartida Actuarial** "
        "y contrastarlo con los requerimientos de la Mesa Actuarial de la Solicitud consolidada "
        "Definitivo v2, incorporando la **integración de la oleada 21–25 de abril de 2026** "
        "(promoción documental y vaciado de **Gerencia Legal** en la mesa). El **detalle contractual de "
        "reaseguro**, el análisis de **TPA en salud** y el **procedimiento administrativo sancionatorio** "
        "asociado a remisiones de contratos de reaseguro se documentan en el **informe confidencial** "
        "`CONF-Accionista-Reaseguro-TPA-Salud-2026-04-24.docx` (distribución restringida), no en esta minuta general.",
        size=11,
    )

    _heading(doc, "2. Alcance cuantitativo y oleadas de trabajo", level=1)

    _heading(doc, "2.1. Carpeta compartida Actuarial (promoción a Mesa)", level=2)
    _p(
        doc,
        "La carpeta reúne alrededor de **176 entregables digitales** en el espejo promovido a "
        "`Reportes/Mesa Actuarial/Materiales/Promocion-2026-04-23-Carpeta-compartida-Actuarial`, "
        "en su mayoría hojas de cálculo Excel, acompañados de informes en PDF y de soportes orientados al "
        "interlocutor regulatorio (SUDEASEG). Los informes de certificación de reservas, reserva complementaria "
        "y ocurridos no reportados constan en el nivel principal de la carpeta compartida actuarial. "
        "La traza documental quedó registrada en `INDICE-PROMOCION-CARPETA-ACTUARIAL-A-MESA.md`.",
    )

    _heading(doc, "2.2. Oleada Semana 2 (21–25 abr. 2026) — cierre informativo", level=2)
    _p(
        doc,
        "En la segunda semana del ciclo POR la mesa consolidó la **lectura integrada** del pliego frente "
        "a lo depositado: (i) publicación de la minuta de validación con código **MIN-POR-ACT-20260424-04** "
        "y archivo Word con **fecha de actualización en el nombre**; (ii) presentación de control "
        "`La_Internacional_Semana2_2026.pptx` ampliada a **7 diapositivas** (síntesis visible en lámina 5; "
        "anexos Legal y Actuarial en láminas 6–7); (iii) **vaciado íntegro** del árbol de "
        "`Reportes/GERENCIA LEGAL` hacia la mesa (176 archivos en espejo y copias temáticas); "
        "(iv) preparación de la **matriz Legal vs relación SUDEASEG 2025** (detalle en informe confidencial); "
        "(v) identificación de riesgos **Alto** en automatización BI y **Medio** en cola de productos SUDEASEG, "
        "analítica distinta de recibos y coordinación SEFA/primas con Finanzas y Tecnología.",
    )

    _heading(doc, "3. Vaciado de Gerencia Legal en Mesa Actuarial", level=1)
    _p(
        doc,
        "Con fecha de espejo **2026-04-23**, se replicó el contenido de `Reportes/GERENCIA LEGAL` en "
        "`Reportes/Mesa Actuarial/Gerencia-Legal/vaciado-completo-2026-04-23/` **sin modificar el origen**. "
        "Para el flujo de trabajo actuarial se clasificaron copias bajo "
        "`Reportes/Mesa Actuarial/Materiales/De-Gerencia-Legal/`, con el mapa "
        "`Reportes/Mesa Actuarial/Gerencia-Legal/INDICE-VACIADO-Y-MAPA-A-MESA-ACTUARIAL.md`. "
        "Estructura temática (resumen):",
    )
    for text in [
        "01-Reaseguro-Contratos — contratos PDF y cruce con SUDEASEG (detalle en informe confidencial).",
        "02-Productos-Condicionados-SUDEASEG — oferta, oficios y relación de productos.",
        "03-Intermediarios-Sucursales; 04-Canales-Alternativos; 05-Plan-Incentivo-esquema-comercial.",
        "06-Sancionatorios-SUDEASEG — documentación sensible (detalle PAS / reaseguro en informe confidencial).",
        "07a-Cumplimiento-AML-FT; 07b-Juicios-activos; 08-Contribuciones-aportes; 09-Solvencias-parafiscales.",
        "10-Presupuesto-y-gobierno-corporativo; 11a-Arrendamientos; 11b-Contratos-proveedores; "
        "12-Declaraciones-ISR; 13-Propiedad-intelectual.",
    ]:
        p = doc.add_paragraph(style="List Bullet")
        r = p.add_run(text)
        _set_run_font(r, size=11)

    _heading(doc, "4. Contenido por bloques", level=1)

    _heading(doc, "4.1. Nivel principal de la carpeta", level=2)
    _p(
        doc,
        "En el nivel principal se concentran los soportes cuantitativos y analíticos del ejercicio "
        "actuarial: modelos y series en hoja de cálculo sobre IBNR, riesgo catastrófico, reaseguro, "
        "relación regulatoria de contratos y movimientos de siniestros al cierre 2025, junto con el "
        "entregable en formato previsto por SUDEASEG para información de recibos de prima.",
    )
    _p(
        doc,
        "Complementan ese bloque los informes en PDF de certificación de reservas, reserva "
        "complementaria, ocurridos no reportados y el procedimiento técnico de siniestralidad "
        "catastrófica. El grado de alineación con la Solicitud consolidada se sintetiza en la "
        "matriz del apartado 6.",
    )

    _heading(doc, "4.2. Baremos convenidos 2026 (APS-ODS)", level=2)
    _p(
        doc,
        "Ciento cincuenta y tres archivos Excel, organizados por entidad federativa y prestador de "
        "salud, con los baremos convenidos entre La Internacional y cada red asistencial para el "
        "año 2026.",
    )

    _heading(doc, "4.3. Margen de solvencia 2025", level=2)
    _p(
        doc,
        "Certificación actuarial del margen de solvencia; formularios MS-01 y MS-02 (incluidas "
        "variantes en PDF); libro de trabajo «MARGEN DE SOLVENCIA DICIEMBRE 2025 LA INTERNACIONAL»; "
        "y publicación en Gaceta Oficial sobre margen de solvencia del sector.",
    )

    _heading(doc, "4.4. Automatización de la reportería de gestión y BI", level=2)
    _p(
        doc,
        "**Riesgo ALTO (gobernanza y continuidad):** no existe hoy un desarrollo ni una tubería automatizada "
        "que extraiga, valide y publique de forma recurrente la reportería de gestión hacia la carpeta Mesa "
        "Actuarial ni hacia los tableros BI (Streamlit / React). La consolidación depende de **depósitos manuales** "
        "y de revisiones puntuales, con coste de tiempo y riesgo de desalineación frente a cierres sucesivos.",
    )
    _p(
        doc,
        "Para el POR, la mesa actuarial debe asumir esta brecha como **dependencia operativa** hasta que "
        "Tecnología y Finanzas definan alcance mínimo de automatización (extracción, control de calidad y "
        "versionado) acotado a los paquetes de cierre que ya se usan en el diagnóstico.",
    )

    _heading(doc, "5. Validación frente a la Solicitud consolidada Definitivo v2", level=1)
    _p(
        doc,
        "La siguiente matriz resume el grado de cobertura de los principales bloques solicitados a "
        "la Mesa Actuarial frente a lo disponible en la carpeta compartida.",
    )
    validation_rows = [
        [
            "Reportería de gestión (cierres y paquetes internos)",
            "Bajo",
            "Paquetes bajo Reportes y copias en Mesa Actuarial",
            "Riesgo operativo bajo pero **proceso manual**; depende de quien depósita y de la coordinación entre áreas.",
        ],
        [
            "Automatización reportería / BI hacia Mesa Actuarial",
            "Alto — sin desarrollo",
            "—",
            "**No hay** extracción ni publicación automatizada hacia la carpeta de trabajo ni hacia los tableros; "
            "prioridad de gobierno del dato para siguientes fases del POR.",
        ],
        [
            "Formularios MS-01 / MS-02 y margen de solvencia",
            "Parcial",
            "Formularios MS, certificación y libro diciembre 2025",
            "Cierre 2025 sólido; confirmar entregas adicionales de 2026 si aplican.",
        ],
        [
            "Certificación de reservas e IBNR",
            "Cubierto",
            "PDF en nivel principal (certificación, ONR, reserva complementaria) y series IBNR",
            "Ubicados en la carpeta base compartida; lectura actuarial posible sin depender de subcarpetas adicionales.",
        ],
        [
            "IBNR — modelos y datos",
            "Cubierto",
            "Series IBNR en Excel (retención y reaseguro)",
            "Base cuantitativa disponible.",
        ],
        [
            "Riesgo catastrófico",
            "Cubierto",
            "Modelo RiesgCAT y procedimiento técnico",
            "Documentación consistente.",
        ],
        [
            "Reaseguro — contratos y relación regulatoria",
            "Parcial — confidencial",
            "Anexo `CONF-Accionista-Reaseguro-TPA-Salud-2026-04-24.docx`",
            "El detalle contractual, matriz por ramo, TPA salud y **PAS SAA-04-1443** se tratan **fuera** de esta minuta general.",
        ],
        [
            "Recibos de prima — formato SUDEASEG",
            "Cubierto",
            "100RPCD2025 (4)",
            "Archivo de intercambio previsto por el regulador para recibos de prima.",
        ],
        [
            "Carga analítica al regulador (distinta de recibos)",
            "Parcial",
            "Recibos de prima cubiertos; otros formatos analíticos",
            "Falta inventariar y depositar **los demás archivos** de carga analítica que no son el archivo de "
            "recibos de prima; coordinar con Tecnología / Finanzas según catálogo SUDEASEG.",
        ],
        [
            "Siniestros pagados y pendientes",
            "Parcial",
            "Series DEF cierre 2025",
            "Homologación fina con plantillas SUDEASEG cuando aplique.",
        ],
        [
            "Baremos por prestador",
            "Cubierto",
            "Baremo referencial y 153 convenios 2026",
            "Cobertura amplia de la red.",
        ],
        [
            "SEFA y comprobantes analíticos (complemento)",
            "Pendiente en carpeta",
            "—",
            "Completar con Finanzas o Tecnología según calendario del POR (complemento a la fila de carga analítica).",
        ],
        [
            "Primas — volcado transaccional",
            "Pendiente en carpeta",
            "—",
            "Coordinar con Operaciones o Tecnología.",
        ],
        [
            "Oferta / cola de productos (trazabilidad SUDEASEG)",
            "Pendiente — en espera",
            "—",
            "A la espera de **confirmación de entrega** de la cola de productos tras **aprobación en SUDEASEG**; "
            "existen condicionados y oficios en `Reportes/CONDICIONADOS DE POLIZAS - PRODUCTOS` y en "
            "`Reportes/GERENCIA LEGAL/CONDICIONADOS DE POLIZAS - PRODUCTOS` para cruce cuando se normalice el flujo.",
        ],
    ]
    avance_pct = _avance_mesa_actuarial_pct(validation_rows)
    _add_table(doc, ["Bloque", "Estado", "Evidencia principal", "Comentario"], validation_rows)

    _heading(doc, "6. Conclusiones", level=1)
    for text in [
        "La carpeta Actuarial aporta un conjunto robusto de certificaciones, reservas, IBNR, "
        "riesgo catastrófico, siniestros, baremos y margen de solvencia, acorde al "
        "espíritu de la Solicitud v2.",
        "La **ausencia de automatización** de reportería y BI es un **riesgo ALTO** de gobierno: conviene "
        "tratarlo como línea de trabajo transversal con Tecnología, no solo como tema de carpeta.",
        "Quedan frentes a completar —carga analítica **más allá de recibos**, SEFA y volcados de primas— "
        "apoyados en las carpetas de **Gerencia Legal** bajo `Reportes` y el vaciado bajo `Mesa Actuarial`. "
        "El **cruce reaseguro** y el **PAS** se gestionan en el **informe confidencial** con Legal.",
        "La Mesa Actuarial puede priorizar con Finanzas y Tecnología el **inventario** de archivos analíticos "
        "pendientes; el trabajo fino de **reaseguro** queda en el anexo restringido.",
    ]:
        p = doc.add_paragraph(style="List Bullet")
        r = p.add_run(text)
        _set_run_font(r, size=11)

    _heading(doc, "7. Próximos pasos", level=1)
    for text in [
        "Acordar con la Lic. Yaritza Oberto el calendario de nuevos depósitos en la carpeta compartida.",
        "Completar la **matriz Legal vs SUDEASEG** en el **informe confidencial** y definir responsable único de mantenimiento.",
        "Definir responsables de los entregables SEFA, de la **carga analítica no-recibos** y de primas pendientes de ubicar.",
        "Canalizar con Tecnología un **alcance mínimo** de automatización de reportería / BI para los cierres que alimentan el POR.",
        "Actualizar el índice ejecutivo de la Mesa Actuarial tras el cierre de homologaciones.",
        "Formalizar respuesta al **SAA-04-1443** solo en el circuito **Legal + informe confidencial** (sin exponer detalle en esta minuta).",
    ]:
        p = doc.add_paragraph(style="List Number")
        r = p.add_run(text)
        _set_run_font(r, size=11)

    _heading(doc, "8. Avance general de la entrega — Mesa Actuarial", level=1)
    p_av = doc.add_paragraph()
    p_av.paragraph_format.space_after = Pt(8)
    r_a0 = p_av.add_run(
        f"Con base en la matriz del apartado 5, el avance global estimado de la información "
        f"disponible en la carpeta compartida Actuarial —respecto de los {len(validation_rows)} bloques contrastados "
        f"con la Solicitud v2— se sitúa en torno al "
    )
    _set_run_font(r_a0, size=11)
    r_a1 = p_av.add_run(f"{avance_pct} %")
    _set_run_font(r_a1, bold=True, size=12, color=TITLE_BLUE)
    r_a2 = p_av.add_run(
        ". Es un indicador sintético (bloques cubiertos, parciales y pendientes en carpeta con "
        "ponderación acordada para seguimiento ejecutivo) y debe actualizarse al cerrar SEFA, "
        "primas y la documentación contractual de reaseguro."
    )
    _set_run_font(r_a2, size=11)

    _heading(
        doc,
        "9. Tabla integrada — Mesa Actuarial (formato Semana 1 + Semana 2, sin detalle reaseguro/PAS)",
        level=1,
    )
    _p(
        doc,
        "La siguiente tabla sigue el **criterio de la lámina «Mesa Actuarial»** de `La_Internacional_Semana1_2026.pdf` "
        "(riesgo en primera columna) y el complemento de la **Semana 2**. Las filas de **reaseguro**, **TPA salud** "
        "y **PAS** se retiraron de la minuta general; figuran en el informe confidencial.",
        size=10,
    )
    filas_hallazgos = _tabla_hallazgos_mesa_sin_detalle_reaseguro()
    _add_table(doc, filas_hallazgos[0], filas_hallazgos[1:])

    doc.add_paragraph()
    foot = doc.add_paragraph()
    fr = foot.add_run(
        f"Confidencial — La Internacional de Seguros — Plan POR 2026. Fecha: {fecha_documento}."
    )
    _set_run_font(fr, italic=True, size=9)

    out_path.parent.mkdir(parents=True, exist_ok=True)
    doc.save(str(out_path))


def main() -> None:
    ap = argparse.ArgumentParser()
    ap.add_argument("--out", type=Path, default=DEFAULT_OUT)
    ap.add_argument("--logo-svg", type=Path, default=DEFAULT_LOGO_SVG)
    ap.add_argument(
        "--logo-png",
        type=Path,
        default=None,
        help="PNG del logo (opcional; por defecto plantillas/logo-la-internacional.png)",
    )
    ap.add_argument("--fecha-doc", default="2026-04-24")
    ap.add_argument("--elaborado-por", default="[Nombre y cargo]")
    ap.add_argument(
        "--referencia-por",
        default=None,
        help="Código POR en tabla «Referencia» (por defecto: MIN-POR-ACT-<AAAAMMDD de --fecha-doc>-03).",
    )
    args = ap.parse_args()
    ref = args.referencia_por
    if not ref:
        compact = args.fecha_doc.replace("-", "")
        ref = f"MIN-POR-ACT-{compact}-04"
    build_document(
        out_path=args.out,
        logo_svg=args.logo_svg,
        logo_png=args.logo_png,
        fecha_documento=args.fecha_doc,
        elaborado_por=args.elaborado_por,
        referencia_por=ref,
    )
    print(f"DOCX generado: {args.out}")


if __name__ == "__main__":
    main()
