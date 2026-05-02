"""
Genera el documento confidencial «Reaseguro y TPA en Salud» con el mismo formato visual
que la minuta Mesa Actuarial (cabecera corporativa, tipografías, tablas).

Incluye el bloque retirado de la minuta general (revisión legal, reaseguro en síntesis, PAS SAA-04-1443),
un **resumen ejecutivo** (menos jerga), **negritas coherentes** (`**texto**` → Word, también en tablas)
y, si se indica fuente, **continúa** el documento con el cuerpo analítico **fusionado** (misma tipografía),
sin apartado meta de «documento aportado».
Opcional: `--matriz-detallada` solo amplía detalle en generaciones internas (no en la versión ejecutiva del informe).

Salidas por defecto:
  - Reportes/Mesa Actuarial/Minutas Reuniones/CONF-Accionista-Reaseguro-TPA-Salud-2026-04-24.docx

Por defecto: minuta **propuesta v1** — síntesis con lectura de **pagos USD** (concentración **Redbridge** y
bloque **no proporcional**); apartado 2 (ranking, **ramos** por contraparte, relación SUDEASEG comprimida);
**Anexo A** (filas **sin fecha** de remisión). ``--fusion-completa`` recupera el flujo antiguo.
"""
from __future__ import annotations

import argparse
import re
import shutil
import unicodedata
from collections import defaultdict
from pathlib import Path
from typing import Any

from docx import Document
from docx.enum.table import WD_ALIGN_VERTICAL, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml.table import CT_Tbl
from docx.oxml.text.paragraph import CT_P
from docx.shared import Inches, Pt, RGBColor
from docx.table import Table
from docx.text.paragraph import Paragraph

import generar_minuta_relevamiento_actuarial_vs_solicitud_docx as minuta_act

TITLE_BLUE = RGBColor(0x1B, 0x3A, 0x5C)
ACCENT_GOLD = RGBColor(0xC9, 0xA2, 0x27)

ROOT_REPORTERIA = Path(__file__).resolve().parents[1]
WORKSPACE = ROOT_REPORTERIA.parent
DEFAULT_LOGO_SVG = WORKSPACE / "La-Internacional-BI" / "public" / "Images" / "Logo horizontal.svg"
DEFAULT_OUT = (
    WORKSPACE
    / "Reportes"
    / "Mesa Actuarial"
    / "Minutas Reuniones"
    / "CONF-Accionista-Reaseguro-TPA-Salud-2026-04-24.docx"
)


def _set_run_font(run, *, bold: bool = False, italic: bool = False, size: int = 11, color=None) -> None:
    run.bold = bold
    run.italic = italic
    run.font.name = "Arial"
    run.font.size = Pt(size)
    run._element.rPr.rFonts.set(qn("w:eastAsia"), "Arial")
    if color is not None:
        run.font.color.rgb = color


def _rule_gold(doc: Document) -> None:
    sep = doc.add_paragraph()
    r = sep.add_run("―" * 62)
    _set_run_font(r, size=8, color=ACCENT_GOLD)


def _iter_block_items(parent):
    for child in parent.element.body.iterchildren():
        if isinstance(child, CT_P):
            yield Paragraph(child, parent)
        elif isinstance(child, CT_Tbl):
            yield Table(child, parent)


def _sintesis_ejecutiva_portada(doc: Document) -> None:
    minuta_act._heading(doc, "Síntesis ejecutiva", level=1)
    for line in [
        "**Reaseguro:** cedencia en **dos monedas**; en **divisas** la **gravedad** viene de la **concentración** en una **sola contraparte** (**Redbridge**) en **salud** y **vida**, con tramos donde puede constar **facultativo al 100%** y administración vía **TPA** (**Third Party Administrator**, tercero administrador de la red) — **alto acople** si esa relación falla.",
        "**Supervisor:** **PAS** (**Procedimiento Administrativo Sancionatorio**) por **información contractual fuera de plazo** (**SAA-04-1443**); riesgo de **sanción** y menor **confianza** del regulador.",
        "**Acción:** la **respuesta** al supervisor debe ir **coordinada** entre **Actuarial**, el **área de reaseguro** y **Legal**; a continuación se desarrolla el **análisis** en **continuidad** con esta síntesis.",
    ]:
        minuta_act._bullet_rich(doc, line, size=11)


def _fold_ascii(s: str) -> str:
    return "".join(
        c for c in unicodedata.normalize("NFD", s.lower()) if unicodedata.category(c) != "Mn"
    )


_RE_RUTAS_ARCH = re.compile(
    r"reportes/|`reportes|[a-z]:\\|vaciado-completo|de-gerencia-legal|indice-vaciado|"
    r"carpeta\s+compartida\s+actuarial|\.md\b|\.xlsx\b|\.pptx\b|\.pdf\b|`[^`]{2,}`",
    re.I,
)

_RE_LISTA_CARPETA = re.compile(
    r"\b\d{2}[a-z]?-[A-Za-záéíóúñÁÉÍÓÚÑ0-9]+(?:-[A-Za-záéíóúñÁÉÍÓÚÑ0-9]+)+\b",
    re.I,
)


def _texto_referencias_o_anexos(s: str) -> bool:
    sl = s.strip().lower()
    if "http://" in sl or "https://" in sl or "www." in sl:
        return True
    if sl.startswith("referencias") or sl == "referencias":
        return True
    if "scribd" in sl or "am best" in sl or "baker mckenzie" in sl:
        return True
    if "accedido el" in sl or "fecha de acceso" in sl:
        return True
    if sl.startswith("confidencial —") and "uso restringido" in sl:
        return True
    if "sudeaseg" in sl and ("gaceta" in sl or "ley de la actividad" in sl):
        return True
    return False


def _bloque_confidencialidad_duplicado(s: str) -> bool:
    f = _fold_ascii(s)
    if "este documento reune informacion sensible" in f and "transfiere riesgos" in f:
        return True
    if "este documento contiene" in f and "informacion sensible" in f and "estructura de reaseguro" in f:
        return True
    if "solo debe circular" in f and "junta directiva" in f:
        return True
    return False


def _linea_bullet_conf_antigua(s: str) -> bool:
    f = _fold_ascii(s)
    keys = (
        "contratos de reaseguro (pdf)",
        "productos y condicionados frente a sudeaseg",
        "intermediarios, canales alternativos",
        "procedimientos sancionatorios",
        "cumplimiento (prevencion de lavado",
        "contratos generales (arrendamiento",
    )
    return any(k in f for k in keys)


def _parrafo_portada_compacta_duplicada(s: str) -> bool:
    f = _fold_ascii(s)
    if len(s) > 500:
        return False
    return "plan de optimizacion y rentabilidad" in f and "informe confidencial" in f


def _parrafo_plantilla_antigua(s: str) -> bool:
    f = _fold_ascii(s)
    if "que es el reaseguro aqui" in f:
        return True
    if "donde esta el tema sensible" in f:
        return True
    if "que paso con el supervisor" in f:
        return True
    if "que sigue en este documento" in f:
        return True
    if "analisis detallado que usted" in f:
        return True
    if "reformateado para que" in f:
        return True
    if "respaldo en carpetas" in f and "tabla resumida" in f:
        return True
    if "para que nadie" in f and "carpetas" in f:
        return True
    if "se copio" in f and "gerencia legal" in f and "mesa actuarial" in f:
        return True
    if "recompuso" in f and "arial" in f:
        return True
    if "sustancia" in f and "reformate" in f:
        return True
    return False


def _texto_operativo_fusion(s: str) -> bool:
    if not s.strip():
        return False
    f = _fold_ascii(s)
    if (
        _texto_referencias_o_anexos(s)
        or _parrafo_plantilla_antigua(s)
        or _bloque_confidencialidad_duplicado(s)
        or _linea_bullet_conf_antigua(s)
        or _parrafo_portada_compacta_duplicada(s)
    ):
        return True
    if _RE_RUTAS_ARCH.search(s):
        return True
    if len(_RE_LISTA_CARPETA.findall(s)) >= 3:
        return True
    needles = (
        "matriz de reaseguro por ramo",
        "extraccion de la relacion",
        "sudeaseg 2025",
        "esquemas contractuales",
        "ejercicio 2025",
        "01-reaseguro",
        "02-productos",
        "06-sancionator",
        "vaciado de gerencia legal",
        "vaciado legal",
        "contenido corporativo y tecnico",
        "indice-vaciado",
        "resumen ejecutivo para el nuevo accionista",
    )
    return any(n in f for n in needles)


def _tabla_detalle_esquemas_regulatorios(rows_data: list[list[str]]) -> bool:
    if not rows_data or not rows_data[0]:
        return False
    h = " ".join(rows_data[0]).lower()
    if "esquemas contractuales" in h:
        return True
    claves = ("tipo de contrato", "capacidad", "retención", "fecha remisión", "reasegurador")
    return sum(1 for k in claves if k in h) >= 4


def _tabla_portada_o_metadatos_duplicados(rows_data: list[list[str]]) -> bool:
    blob = " | ".join(" ".join(r) for r in rows_data[:4])
    return "por 2026" in blob.lower() and "informe confidencial" in blob.lower() and len(rows_data) <= 5


_RE_INLINE_URL = re.compile(r"https?://\S+|www\.\S+", re.I)
_RE_BRACKET_CITE = re.compile(r"\[\s*\d+\s*\]|\[[0-9,\s]{1,24}\]")
_RE_PAREN_YEAR = re.compile(r"\(\s*[12][0-9]{3}\s*\)")

_KEYWORDS_SCORE: tuple[tuple[str, int], ...] = (
    ("reaseguro", 3),
    ("cedencia", 2),
    ("contraparte", 2),
    ("redbridge", 3),
    ("facultativo", 2),
    ("proporcional", 2),
    ("tratado", 2),
    ("capa", 2),
    ("prima", 2),
    ("tpa", 3),
    ("tercer", 2),
    ("administrador", 2),
    ("salud", 2),
    ("pas", 3),
    ("sancionator", 3),
    ("supervisor", 2),
    ("sudeaseg", 2),
    ("saa-04", 3),
    ("saa 04", 3),
    ("regulador", 2),
    ("riesgo", 2),
    ("concentrac", 2),
    ("legal", 2),
    ("actuarial", 2),
)


def _omit_paragraph_fuente_informe(t: str) -> bool:
    if not t.strip():
        return True
    sl = t.strip().lower()
    f = _fold_ascii(t)
    if len(sl) < 4:
        return True
    if "scribd" in sl or "am best" in sl or "baker mckenzie" in sl:
        return True
    if "accedido el" in sl or "fecha de acceso" in sl or "recuperado el" in sl:
        return True
    if sl.startswith("confidencial —") and "uso restringido" in sl:
        return True
    if f.startswith("nota metodologica"):
        return True
    if "http://" in sl or "https://" in sl:
        sin_url = _RE_INLINE_URL.sub("", sl)
        if len(sin_url.strip()) < max(30, len(sl) // 4):
            return True
    return False


def _strip_citas_inline(s: str) -> str:
    s = _RE_INLINE_URL.sub("", s)
    s = re.sub(r"\bdoi:\s*\S+", "", s, flags=re.I)
    s = _RE_BRACKET_CITE.sub("", s)
    s = _RE_PAREN_YEAR.sub("", s)
    s = re.sub(r"\s+", " ", s).strip(" -—\t")
    return s


def _collect_text_from_source_docx(path: Path) -> str:
    if not path.is_file():
        return ""
    src = Document(str(path))
    parts: list[str] = []
    drop_referencias = False
    for block in _iter_block_items(src):
        if isinstance(block, Paragraph):
            t = " ".join(block.text.split())
            if not t:
                continue
            line0 = t.splitlines()[0].strip()
            fo = _fold_ascii(line0)
            if fo.startswith("referencias") or fo.startswith("bibliografia"):
                drop_referencias = True
            if drop_referencias:
                continue
            if _omit_paragraph_fuente_informe(t):
                continue
            if _parrafo_portada_compacta_duplicada(t):
                continue
            cleaned = _strip_citas_inline(t)
            if len(cleaned) > 8:
                parts.append(cleaned)
        elif isinstance(block, Table):
            for row in block.rows:
                for cell in row.cells:
                    t = " ".join(cell.text.split())
                    if not t or drop_referencias:
                        continue
                    if _omit_paragraph_fuente_informe(t):
                        continue
                    cleaned = _strip_citas_inline(t)
                    if len(cleaned) > 8:
                        parts.append(cleaned)
    return "\n".join(parts)


def _split_sentences_es(blob: str) -> list[str]:
    blob = blob.replace("\n", " ")
    blob = re.sub(r"\s+", " ", blob).strip()
    if not blob:
        return []
    raw = re.split(r"(?<=[\.\!\?])\s+", blob)
    out: list[str] = []
    for p in raw:
        p = p.strip()
        if len(p) > 18:
            out.append(p)
    return out


def _sentence_score(s: str) -> int:
    f = _fold_ascii(s)
    n = 0
    for kw, w in _KEYWORDS_SCORE:
        if kw in f:
            n += w
    if re.search(r"\d", s):
        n += 1
    return n


def _compress_sentences(sentences: list[str], ratio: float) -> list[str]:
    seen: set[str] = set()
    uniq: list[str] = []
    for s in sentences:
        k = _fold_ascii(s[:160])
        if k in seen:
            continue
        seen.add(k)
        uniq.append(s)
    total = sum(len(s) + 1 for s in uniq)
    if not uniq:
        return []
    r = max(0.18, min(0.55, ratio))
    max_chars = max(900, int(total * r))
    out: list[str] = []
    used = 0
    for s in uniq:
        sc = _sentence_score(s)
        prospective = used + len(s) + 1
        if prospective > max_chars:
            if sc < 2:
                continue
            prospective = used + len(s) + 1
            if prospective > max_chars:
                break
        out.append(s)
        used = prospective
    return out


def _paragraphs_from_sentences(sents: list[str], *, chars_per_para: int = 520) -> list[str]:
    paras: list[str] = []
    buf: list[str] = []
    n = 0
    for s in sents:
        add_len = len(s) + (1 if buf else 0)
        if buf and n + add_len > chars_per_para:
            paras.append(" ".join(buf))
            buf = [s]
            n = len(s) + 1
        else:
            if buf:
                n += 1
            buf.append(s)
            n += len(s)
    if buf:
        paras.append(" ".join(buf))
    return paras


def _try_float_cell(v) -> float | None:
    if v is None:
        return None
    if isinstance(v, (int, float)):
        return float(v)
    s = str(v).strip().replace(",", "").replace(" ", "")
    if not s:
        return None
    try:
        return float(s)
    except ValueError:
        return None


def _normalize_empresa_pagos(label: str) -> str:
    s = re.sub(r"\s+", " ", str(label)).strip()
    s = re.sub(r"^[\d.,\s%]+$", "", s).strip()
    if not s:
        return "—"
    low = s.lower()
    if "redbridge" in low or "redbrige" in low:
        return "Redbridge"
    if low.startswith("kairos"):
        return "Kairos"
    if low.startswith("bee"):
        return "Bee"
    if "hannover" in low and "kairos" in low:
        return "Hannover / Kairos (reparto declarado)"
    if "r.i.v" in low or low.startswith("riv"):
        return "R.I.V"
    if "lockton" in low:
        return "Lockton"
    if "worldwide" in low:
        return "Worldwide"
    if "delta" in low:
        return "Delta"
    return s[:48]


def _extract_ramo_desde_etiqueta_pagos(raw: str, emp: str) -> str | None:
    t = re.sub(r"\s+", " ", str(raw)).strip()
    if not t:
        return None
    t2 = re.sub(rf"^{re.escape(emp.strip())}\s*:?\s*", "", t, flags=re.I).strip()
    if len(t2) < 5:
        return None
    return t2[:160]


def _parse_pagos_usd_excel(path: Path) -> dict[str, Any] | None:
    """
    Ranking USD, total de control, suma no proporcional atribuible a Redbridge y ramos por contraparte
    (según etiquetas y bloques del Excel de pagos).
    """
    try:
        import openpyxl
    except ImportError:
        return None
    if not path.is_file():
        return None
    wb = openpyxl.load_workbook(path, read_only=True, data_only=True)
    ws = wb[wb.sheetnames[0]]
    en_dolares = False
    bucket = "Pagos en USD"
    acc: dict[str, dict] = defaultdict(lambda: {"amt": 0.0, "lines": [], "ramos": set()})
    control_total: float | None = None
    np_redbridge = 0.0
    for row in ws.iter_rows(values_only=True):
        if not row:
            continue
        c1 = row[1] if len(row) > 1 else None
        c2 = row[2] if len(row) > 2 else None
        s1 = str(c1).strip() if c1 is not None else ""
        u1 = s1.upper()
        if "PRODUCTOS EN BOLIV" in u1 or ("BOLIV" in u1 and "PRODUCTOS" in u1):
            en_dolares = False
            continue
        if "PRODUCTOS EN DOL" in u1 or ("DOL" in u1 and "PRODUCTOS" in u1):
            en_dolares = True
            continue
        if not en_dolares and "BOLIV" not in u1 and "DOL" in u1 and ("REASEGUROS" in u1 or "PRODUCTOS" in u1):
            en_dolares = True
        if "TOTAL PAGOS EN USD" in u1:
            control_total = _try_float_cell(c2)
            continue
        if not en_dolares:
            continue
        ul = s1.lower()
        if "bouquet" in ul and "proporcional" in ul and "usd" in ul:
            bucket = "Bouquet proporcional (USD)"
            continue
        if "no proporcional" in ul and "usd" in ul:
            bucket = "No proporcional (USD)"
            continue
        if "facultativo" in ul and "puro" in ul and "usd" in ul:
            bucket = "Facultativo puro (USD)"
            continue
        amt = _try_float_cell(c2)
        if amt is None or not s1:
            continue
        if "total" in ul and "prop" in ul:
            continue
        if re.fullmatch(r"[\d.,\s]+", s1):
            continue
        key = _normalize_empresa_pagos(s1)
        acc[key]["amt"] += amt
        desc = s1[:72]
        if len(acc[key]["lines"]) < 8 and desc not in acc[key]["lines"]:
            acc[key]["lines"].append(desc)
        frag = _extract_ramo_desde_etiqueta_pagos(s1, key)
        if frag:
            acc[key]["ramos"].add(f"{frag} — {bucket}")
        else:
            acc[key]["ramos"].add(bucket)
        if bucket == "No proporcional (USD)" and key == "Redbridge":
            np_redbridge += amt
    wb.close()
    tot_sum = sum(d["amt"] for d in acc.values()) or 0.0
    total_usd = float(control_total) if control_total is not None else float(tot_sum or 0.0)
    ranked = sorted(acc.items(), key=lambda kv: -kv[1]["amt"])
    out: list[tuple[str, float, list[str]]] = []
    ramos_por: dict[str, list[str]] = {}
    rb_total = 0.0
    for k, d in ranked:
        if k == "—" or d["amt"] <= 0:
            continue
        out.append((k, d["amt"], d["lines"]))
        ramos_por[k] = sorted(d["ramos"])
        if k == "Redbridge":
            rb_total = float(d["amt"])
    return {
        "ranking": out,
        "total_usd": total_usd,
        "np_redbridge_usd": np_redbridge,
        "redbridge_total_usd": rb_total,
        "ramos_por_empresa": ramos_por,
    }


def _bullet_sintesis_reaseguro_flujos(parsed: dict[str, Any] | None) -> str:
    if not parsed or not parsed.get("total_usd"):
        return (
            "**Dinero y contrapartes:** la compañía mueve reaseguro en **dólares** y en **bolívares**; en dólares casi todo el volumen "
            "va a **muy pocas** aseguradoras reaseguradoras (el detalle numérico está en el apartado 2)."
        )
    tot = float(parsed["total_usd"])
    rb = float(parsed.get("redbridge_total_usd") or 0.0)
    np_rb = float(parsed.get("np_redbridge_usd") or 0.0)
    if tot <= 0 or rb <= 0:
        return (
            "**Dinero y contrapartes:** la compañía mueve reaseguro en **dólares** y en **bolívares**; en dólares casi todo el volumen "
            "va a **muy pocas** aseguradoras reaseguradoras (el detalle numérico está en el apartado 2)."
        )
    pct_rb = 100.0 * rb / tot
    pct_np_tot = 100.0 * np_rb / tot if tot else 0.0
    pct_np_rb = 100.0 * np_rb / rb if rb else 0.0
    return (
        f"**Dinero y contrapartes:** casi **{pct_rb:.0f} de cada 100 dólares** que la compañía paga por reaseguro van a **Redbridge** "
        f"(unos **{rb:,.0f} USD** de alrededor de **{tot:,.0f} USD** en total). Dentro de eso, solo el tramo de **salud** que figura como "
        f"**no proporcional** —planes internacionales y salud «técnica» tradicional— suma unos **{np_rb:,.0f} USD**, o sea **{pct_np_tot:.0f} %** "
        f"de todo lo pagado en dólares y **{pct_np_rb:.0f} %** de lo que se le paga a **Redbridge**; por eso cualquier fallo ahí **duele** en la cuenta y en la reputación."
    )


def _bullet_sintesis_gravedad_regulador(parsed: dict[str, Any] | None) -> str:
    base_np = float(parsed.get("np_redbridge_usd") or 0.0) if parsed else 0.0
    if base_np > 0:
        millones = base_np / 1_000_000
        return (
            f"**Por qué es tan delicado:** ese bloque de **salud** agrupa **muchísimo dinero** (orden de **{millones:.1f} millones de dólares**). "
            "Que la información haya llegado **tarde** a **SUDEASEG** (superintendencia de seguros) y que, **hoy por hoy**, **no** esté a la vista el "
            "**contrato** que ampare ese negocio **no es un detalle**: mezcla riesgo de **multa o proceso sancionatorio**, **mala imagen** ante el regulador "
            "y **preguntas incómodas** de accionistas o auditores."
        )
    return (
        "**Por qué es delicado:** informar **tarde** al regulador sobre contratos de reaseguro **grandes**, sin tener claro el **contrato** que los respalda, "
        "expone a la compañía en varios frentes a la vez."
    )


def _bullet_sintesis_regulador_simple() -> str:
    return (
        "**Lista al regulador:** lo que la compañía declara cada año a **SUDEASEG** debe **cuadrar** con lo que realmente se paga; "
        "si hay filas **sin fecha** de envío, el detalle está al final en el **Anexo A** para no alargar esta lectura."
    )


def _bullet_sintesis_supervisor_simple() -> str:
    return (
        "**Trámite con el supervisor:** está abierto un **PAS** (**Procedimiento Administrativo Sancionatorio**) —el canal del supervisor para **investigar** "
        "y eventualmente **sancionar**— con la referencia **SAA-04-1443**, porque parte de la **documentación de reaseguro** se envió **fuera de plazo**. "
        "La respuesta debe salir **ordenada** entre quienes entienden los **números** (**Actuarial** y **reaseguro**) y **Legal**."
    )


def _p_confidencialidad_simple(doc: Document) -> None:
    minuta_act._p_rich(
        doc,
        "**Confidencial:** este texto es solo para **accionistas de confianza**, **dirección** y los equipos de **Legal** y **Actuarial**; "
        "**no** sirve para circulares públicas del **POR**.",
        size=10,
    )


def _sudeaseg_sheet_incluir(sname: str) -> bool:
    sn = sname.strip()
    if sn == "Hoja1":
        return True
    if "2023" in sn and "2025" not in sn:
        return False
    return "2025" in sn


def _sudeaseg_fecha_remision_vacia(rem: str) -> bool:
    if not rem or not str(rem).strip():
        return True
    r = str(rem).strip().lower()
    if r in ("n/a", "-", "—", "na"):
        return True
    return False


def _iter_filas_contrato_sudeaseg(path: Path) -> tuple[list[dict], str]:
    """Filas de contrato con moneda, tipo, ramo, reaseguradores, fecha remisión."""
    try:
        import openpyxl
    except ImportError:
        return [], "openpyxl no disponible."
    if not path.is_file():
        return [], "No se localizó el Excel de relación SUDEASEG."
    out: list[dict] = []
    wb = openpyxl.load_workbook(path, read_only=True, data_only=True)
    for sname in wb.sheetnames:
        if not _sudeaseg_sheet_incluir(sname):
            continue
        moneda = "VES" if "Bs" in sname else "USD"
        ws = wb[sname]
        start = None
        for row in ws.iter_rows(values_only=True):
            if row and str(row[0]).strip() == "Tipo de Contrato":
                start = True
                continue
            if not start:
                continue
            if not row or row[0] is None:
                continue
            tipo = str(row[0]).strip()
            if not tipo:
                continue
            ramo = str(row[4]).strip()[:80] if len(row) > 4 and row[4] else ""
            reas_cel = row[11] if len(row) > 11 else None
            rem = row[14] if len(row) > 14 else None
            nombres = minuta_act._reas_nombres_desde_celda(reas_cel)
            if not nombres:
                nombres = ["(celda contraparte vacía o no legible)"]
            out.append(
                {
                    "hoja": sname.strip(),
                    "moneda": moneda,
                    "tipo": tipo[:40],
                    "ramo": ramo,
                    "nombres": nombres,
                    "rem": rem,
                    "sin_rem": _sudeaseg_fecha_remision_vacia(str(rem) if rem is not None else ""),
                }
            )
    wb.close()
    return out, path.name


def _agrupa_regulador_comprimido(filas: list[dict]) -> list[list[str]]:
    """Contraparte | contratos | remitidos | pendientes | ramos (muestra). Una fila de relación → una contraparte principal."""
    g: dict[str, dict] = defaultdict(lambda: {"n": 0, "ok": 0, "pend": 0, "ramos": set(), "label": ""})
    for r in filas:
        nm0 = r["nombres"][0] if r["nombres"] else "(sin contraparte)"
        key = minuta_act._fold_key(nm0[:52])
        if not g[key]["label"]:
            g[key]["label"] = nm0.strip()[:56]
        g[key]["n"] += 1
        if r["sin_rem"]:
            g[key]["pend"] += 1
        else:
            g[key]["ok"] += 1
        if r["ramo"]:
            g[key]["ramos"].add(r["ramo"][:42])
    rows: list[list[str]] = []
    for _, d in sorted(g.items(), key=lambda kv: -(kv[1]["n"]))[:12]:
        muestra = ", ".join(sorted(d["ramos"])[:4])
        if len(d["ramos"]) > 4:
            muestra += "…"
        rows.append(
            [
                d.get("label", "—"),
                str(d["n"]),
                str(d["ok"]),
                str(d["pend"]),
                muestra or "—",
            ]
        )
    return rows


def _filas_sin_remision_tabla(filas: list[dict], *, max_rows: int = 16) -> list[list[str]]:
    cand = [r for r in filas if r["sin_rem"]]
    cand.sort(key=lambda x: (0 if x["moneda"] == "USD" else 1, x["tipo"], x["ramo"]))
    rows: list[list[str]] = []
    for r in cand:
        corte = " · ".join(r["nombres"][:2])[:90]
        mark = "**Sí** — **sin fecha** de remisión en la relación"
        rows.append([r["moneda"], r["tipo"], r["ramo"][:56], corte, mark])
        if len(rows) >= max_rows:
            break
    return rows


def _append_anexo_a_sin_remision_sudeaseg(doc: Document) -> None:
    p_sude = minuta_act._find_sudeaseg_reaseguro_xlsx()
    if not p_sude or not p_sude.is_file():
        return
    filas_sude, nombre_sude = _iter_filas_contrato_sudeaseg(p_sude)
    sin_r = _filas_sin_remision_tabla(filas_sude)
    if not sin_r:
        return
    minuta_act._heading(doc, "Anexo A. Detalle — relación SUDEASEG sin fecha de remisión", level=1)
    minuta_act._p(
        doc,
        f"Fuente: **{nombre_sude}**. Uso interno para cierre con **Legal** y **reaseguro**.",
        size=9,
        italic=True,
    )
    minuta_act._p_rich(
        doc,
        "**Importante:** estas filas aparecen en la **lista al regulador** pero **sin fecha** de haber sido **enviadas** a **SUDEASEG**; "
        "conviene **cerrarlas** pronto con **Legal** y **reaseguro**.",
        size=11,
    )
    minuta_act._add_table(
        doc,
        ["Moneda", "Tipo", "Ramo", "Contraparte", "Estado remisión"],
        sin_r,
    )


def _append_propuesta_v1_bloques_datos(doc: Document) -> None:
    p_pagos = minuta_act._find_resumen_pagos_reaseguro_xlsx()
    parsed = _parse_pagos_usd_excel(p_pagos) if p_pagos and p_pagos.is_file() else None

    minuta_act._heading(doc, "1. Síntesis para el accionista", level=1)
    minuta_act._bullet_rich(doc, _bullet_sintesis_reaseguro_flujos(parsed), size=11)
    minuta_act._bullet_rich(doc, _bullet_sintesis_gravedad_regulador(parsed), size=11)
    minuta_act._bullet_rich(doc, _bullet_sintesis_regulador_simple(), size=11)
    minuta_act._bullet_rich(doc, _bullet_sintesis_supervisor_simple(), size=11)
    _p_confidencialidad_simple(doc)

    minuta_act._heading(
        doc,
        "2. Cifras en dólares, ranking y lo declarado al regulador",
        level=1,
    )
    minuta_act._p_rich(
        doc,
        "Aquí se cruza, en **lenguaje sencillo**, lo que la compañía **paga** por reaseguro en **dólares** con lo que **figura** en la **lista anual** "
        "que se envía a **SUDEASEG**; no es una lectura actuarial fina, sino una **foto** para la **junta**.",
        size=11,
    )

    p_sude = minuta_act._find_sudeaseg_reaseguro_xlsx()
    filas_sude, nombre_sude = _iter_filas_contrato_sudeaseg(p_sude) if p_sude else ([], "")

    if parsed and p_pagos and p_pagos.is_file():
        ranking = parsed["ranking"]
        base = float(parsed["total_usd"]) or sum(x[1] for x in ranking) or 1.0
        np_rb = float(parsed.get("np_redbridge_usd") or 0.0)
        minuta_act._heading(doc, "2.1. Quién recibe los dólares (ranking interno)", level=2)
        minuta_act._p_rich(
            doc,
            f"Cifras tomadas del archivo interno **{p_pagos.name}**. El total en **dólares** que aparece en ese cuadro es de **≈ {base:,.0f} USD**.",
            size=10,
        )
        tbl_rows: list[list[str]] = []
        for emp, amt, _lines in ranking[:14]:
            pct = 100.0 * amt / base
            tbl_rows.append([emp, f"≈ {amt:,.0f} USD", f"≈ {pct:.1f} %"])
        if tbl_rows:
            minuta_act._add_table(doc, ["Contraparte (nombre resumido)", "Dólares pagados (aprox.)", "Parte del total"], tbl_rows)
        if np_rb > 0:
            minuta_act._p_rich(
                doc,
                f"**Lectura clave:** buena parte de ese dinero está en un **solo tipo** de operación con **Redbridge** (salud **no proporcional**, "
                f"unos **{np_rb:,.0f} USD**). Ahí es donde duele más que el **papel** del contrato **no** esté **cerrado** y que la **comunicación** al regulador haya sido **extemporánea**.",
                size=11,
            )
    else:
        minuta_act._p_rich(doc, "**Aviso:** no se encontró el Excel de **resumen de pagos** en `Reportes`.", size=10)

    if filas_sude:
        minuta_act._heading(doc, "2.2. Lo que se le dice al regulador (resumen por contraparte)", level=2)
        minuta_act._p_rich(
            doc,
            "Aquí se resume la **información que la compañía declara** a partir del **detalle de todos los contratos** de reaseguro que van en la **lista anual** "
            "a **SUDEASEG**: cada **fila** del archivo es un contrato o tramo (tipo, ramo, capacidad, contrapartes, participación, etc.). "
            "En este documento **no** se repite contrato por contrato: se **agrupa** por **contraparte** para una lectura rápida; el detalle línea a línea está en el archivo "
            f"**{nombre_sude}** (no sustituye leer el **PDF** del contrato cuando haga falta).",
            size=10,
        )
        comp = _agrupa_regulador_comprimido(filas_sude)
        if comp:
            minuta_act._add_table(
                doc,
                ["Contraparte (como en relación)", "Líneas", "Con fecha remisión", "Sin fecha remisión", "Ramos (muestra)"],
                comp,
            )
        sin_r = _filas_sin_remision_tabla(filas_sude)
        if sin_r:
            minuta_act._p_rich(
                doc,
                "Existen líneas en la relación **sin fecha** de remisión a SUDEASEG: el **detalle tabular** se traslada al **Anexo A** "
                "para no saturar el cuerpo principal.",
                size=10,
            )
        else:
            minuta_act._p(doc, "En el extracto automático no quedaron filas **sin fecha** de remisión.", size=10, italic=True)
    else:
        minuta_act._p_rich(doc, "**Aviso:** no se pudo leer la **relación SUDEASEG 2025** para la vista comprimida.", size=10)

    minuta_act._heading(doc, "3. Supervisor — PAS (SAA-04-1443)", level=1)
    minuta_act._p_rich(
        doc,
        "El **PAS** (**Procedimiento Administrativo Sancionatorio**) es el trámite formal por el que el **supervisor** puede **investigar** incumplimientos y, "
        "si procede, **sancionar** a la compañía. En este caso hay una **carta oficial** (**SAA-04-1443**) que **abre** ese procedimiento porque entiende que "
        "la compañía **no** entregó a tiempo parte de la **documentación contractual de reaseguro**. Eso puede acabar en **multas** u otras medidas **graves**, "
        "y en todo caso **perjudica** la confianza con el regulador. La respuesta debe ser **una sola voz** entre quienes dominan los **números** (**Actuarial** y **reaseguro**) "
        "y **Legal**.",
        size=11,
    )
    minuta_act._p_rich(
        doc,
        "**Enlace con el punto anterior:** cuando el **dinero** en juego es **muy alto** y además **no** está claro el **contrato** que lo respalda, un **PAS** "
        "(**Procedimiento Administrativo Sancionatorio**) se percibe **aún más delicado**; no es solo un retraso de papeles, es **credibilidad** frente a **SUDEASEG**.",
        size=11,
    )
    minuta_act._heading(doc, "4. TPA en salud y esquema tipo «fronting» (hipótesis)", level=1)
    minuta_act._p_rich(
        doc,
        "En **salud** suele intervenir un **tercero** que **administra** la red y los **pagos** de siniestros; en el sector eso se llama **TPA** "
        "(**Third Party Administrator**). **Fronting** es una palabra fea pero útil: describe el caso en que **La Internacional** queda como **cara visible** "
        "del seguro frente al cliente, pero el **riesgo de caja** y el **día a día** dependen mucho de **otros contratos** y del **reaseguro** que hay detrás.",
        size=11,
    )
    minuta_act._p_rich(
        doc,
        "**Hipótesis (hay que profundizar):** hace falta **más trabajo** con **Legal** y con quienes conocen la **operación** para **confirmar** o **descartar** "
        "este esquema. Con lo que **consta** en los **papeles** revisados hasta ahora, la **cedencia** aparece como **ciento por ciento** (**100 %**); cuando el "
        "**reaseguro** cubre **todo** así, lo **más natural** de leer —aunque **no** sea la **única** explicación— es un modelo donde un **tercero** (**TPA**) "
        "lleva la **operación** y la aseguradora actúa más como **canal**.",
        size=11,
    )
    minuta_act._p_rich(
        doc,
        "**Qué conviene vigilar:** que lo que **dice** la póliza al público, lo que **ordena** el tercero administrador y lo que **cubre** el reaseguro "
        "**cuadre**; mientras haya **dudas**, conviene **no** crecer en volumen hasta que **Legal** y **reaseguro** cierren la lectura.",
        size=11,
    )


def _build_header_minuta_accionista(
    doc: Document,
    *,
    fecha_documento: str,
    referencia: str,
    elaborado_por: str,
) -> None:
    style = doc.styles["Normal"]
    style.font.name = "Arial"
    style.font.size = Pt(11)
    style._element.rPr.rFonts.set(qn("w:eastAsia"), "Arial")

    pic_path, pic_is_temp = minuta_act._resolve_logo_path(logo_svg=DEFAULT_LOGO_SVG, logo_png=None)
    logo_stream = minuta_act._logo_png_bytesio(pic_path) if pic_path is not None else None

    enc = doc.add_table(rows=1, cols=2)
    enc.autofit = False
    enc.alignment = WD_TABLE_ALIGNMENT.CENTER
    w_logo_col = Inches(3.05)
    w_txt_col = Inches(3.15)
    enc.columns[0].width = w_logo_col
    enc.columns[1].width = w_txt_col
    c_logo = enc.rows[0].cells[0]
    c_txt = enc.rows[0].cells[1]
    c_logo.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
    c_txt.vertical_alignment = WD_ALIGN_VERTICAL.CENTER

    p_logo = c_logo.paragraphs[0]
    p_logo.alignment = WD_ALIGN_PARAGRAPH.LEFT
    p_logo.paragraph_format.space_before = Pt(0)
    p_logo.paragraph_format.space_after = Pt(0)
    if logo_stream is not None:
        p_logo.add_run().add_picture(logo_stream, width=Inches(2.75))
    elif pic_path is not None and pic_path.is_file():
        p_logo.add_run().add_picture(str(pic_path), width=Inches(2.75))
    else:
        r = p_logo.add_run("La Internacional\nde Seguros")
        _set_run_font(r, bold=True, size=8, color=TITLE_BLUE)

    if pic_is_temp and pic_path is not None:
        try:
            pic_path.unlink(missing_ok=True)
        except OSError:
            pass

    pr = c_txt.paragraphs[0]
    pr.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    pr.paragraph_format.space_before = Pt(0)
    pr.paragraph_format.space_after = Pt(0)
    pr.paragraph_format.right_indent = Pt(8)
    r1 = pr.add_run("La Internacional de Seguros, C.A.\n")
    _set_run_font(r1, bold=True, size=9, color=TITLE_BLUE)
    r2 = pr.add_run("Plan de Optimización y Rentabilidad (POR) 2026\n")
    _set_run_font(r2, size=8)
    r3 = pr.add_run("Mesa Actuarial · Lectura para accionista")
    _set_run_font(r3, italic=True, size=8, color=ACCENT_GOLD)

    doc.add_paragraph()
    sep = doc.add_paragraph()
    run_sep = sep.add_run("―" * 62)
    _set_run_font(run_sep, size=8, color=ACCENT_GOLD)

    t = doc.add_paragraph()
    t.alignment = WD_ALIGN_PARAGRAPH.CENTER
    tr = t.add_run("MINUTA — LECTURA ACCIONISTA\n")
    _set_run_font(tr, bold=True, size=18, color=TITLE_BLUE)
    tr2 = t.add_run("Reaseguro, TPA en salud y supervisión (versión resumida)")
    _set_run_font(tr2, bold=True, size=12, color=TITLE_BLUE)

    doc.add_paragraph()
    minuta_act._add_table(
        doc,
        ["Campo", "Detalle"],
        [
            ["Referencia", referencia],
            ["Fecha del documento", fecha_documento],
            [
                "Insumos de trabajo",
                "**Control de pagos de reaseguro (Excel)** y **relación contractual SUDEASEG 2025**; circulación **acotada**.",
            ],
            [
                "Asunto",
                "Lectura para **accionista de referencia** y **alta dirección** sobre **reaseguro**, **TPA** en salud y **PAS**.",
            ],
            ["Clasificación", "Confidencial — circulación acotada"],
            ["Elaborado por", elaborado_por],
        ],
    )


def _append_user_analysis_integrated(target: Document, source_path: Path) -> None:
    """Fusiona el cuerpo analítico recorriendo bloques, con tipografía homogénea (Arial 11), sin encabezado meta."""
    if not source_path.is_file():
        return
    src = Document(str(source_path))
    _dup_subtitulo = (
        "Análisis de reaseguro, TPA en salud y procedimiento administrativo (PAS)".casefold()
    )

    def _omitir_titulo_fusion(s: str) -> bool:
        raw = s.strip()
        if not raw:
            return False
        line0 = " ".join(raw.splitlines()[0].split())
        f0 = _fold_ascii(line0)
        if f0 == "informe confidencial" or f0.startswith("informe confidencial "):
            return True
        if " ".join(raw.split()).casefold() == _dup_subtitulo:
            return True
        if "contenido corporativo" in f0 and "tecnico" in f0:
            return True
        if "vaciado" in f0 and "gerencia legal" in f0:
            return True
        if "mesa actuarial" in f0 and "trazabilidad" in f0:
            return True
        if f0 == _fold_ascii("Confidencialidad y uso interno"):
            return True
        if "resumen ejecutivo" in f0 and ("accionista" in f0 or "nuevo" in f0):
            return True
        if "esquemas contractuales" in f0:
            return True
        if "reaseguro" in f0 and "esquema" in f0 and "sudeaseg" in f0:
            return True
        if "reaseguro" in f0 and "matriz" in f0 and "ramo" in f0:
            return True
        if "antecedentes" in f0 and "respaldo" in f0 and "documental" in f0:
            return True
        if "analisis de profundidad" in f0:
            return True
        if "documento base integrado" in f0:
            return True
        if f0 == _fold_ascii("Síntesis ejecutiva") or f0 == _fold_ascii("Sintesis ejecutiva"):
            return True
        if f0 == _fold_ascii("Resumen ejecutivo"):
            return True
        if "informe de auditoria integral" in f0 and "estructura de reaseguro" in f0:
            return True
        return False

    def _titulo_fusion_limpio(t: str) -> str:
        return re.sub(r"^\s*\d+(\.\d+)*\s*[.—\-:]*\s*", "", t.strip())

    for block in _iter_block_items(src):
        if isinstance(block, Paragraph):
            t = block.text.strip()
            if not t:
                target.add_paragraph()
                continue
            if _omitir_titulo_fusion(t):
                continue
            if _texto_operativo_fusion(t):
                continue
            st = block.style.name if block.style else "Normal"
            if st == "Title" or st == "Heading 1":
                minuta_act._heading(target, _titulo_fusion_limpio(t), level=2)
            elif st.startswith("Heading"):
                minuta_act._heading(target, _titulo_fusion_limpio(t), level=3)
            else:
                p = target.add_paragraph()
                p.paragraph_format.space_after = Pt(6)
                if block.runs:
                    for run in block.runs:
                        if run.text == "":
                            continue
                        rr = p.add_run(run.text)
                        _set_run_font(rr, bold=bool(run.bold), italic=bool(run.italic), size=11)
                else:
                    rr = p.add_run(t)
                    _set_run_font(rr, size=11)
        elif isinstance(block, Table):
            rows_data: list[list[str]] = []
            for row in block.rows:
                rows_data.append([cell.text.replace("\n", " ").strip() for cell in row.cells])
            if not rows_data:
                continue
            blob = " | ".join(" ".join(r) for r in rows_data)
            if (
                _texto_operativo_fusion(blob)
                or _tabla_detalle_esquemas_regulatorios(rows_data)
                or _tabla_portada_o_metadatos_duplicados(rows_data)
            ):
                continue
            hdr = rows_data[0]
            body_rows = rows_data[1:] if len(rows_data) > 1 else []
            minuta_act._add_table(target, hdr, body_rows)


def _build_header_confidencial(
    doc: Document,
    *,
    fecha_documento: str,
    referencia: str,
    elaborado_por: str,
) -> None:
    style = doc.styles["Normal"]
    style.font.name = "Arial"
    style.font.size = Pt(11)
    style._element.rPr.rFonts.set(qn("w:eastAsia"), "Arial")

    pic_path, pic_is_temp = minuta_act._resolve_logo_path(logo_svg=DEFAULT_LOGO_SVG, logo_png=None)
    logo_stream = minuta_act._logo_png_bytesio(pic_path) if pic_path is not None else None

    enc = doc.add_table(rows=1, cols=2)
    enc.autofit = False
    enc.alignment = WD_TABLE_ALIGNMENT.CENTER
    w_logo_col = Inches(3.05)
    w_txt_col = Inches(3.15)
    enc.columns[0].width = w_logo_col
    enc.columns[1].width = w_txt_col
    c_logo = enc.rows[0].cells[0]
    c_txt = enc.rows[0].cells[1]
    c_logo.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
    c_txt.vertical_alignment = WD_ALIGN_VERTICAL.CENTER

    p_logo = c_logo.paragraphs[0]
    p_logo.alignment = WD_ALIGN_PARAGRAPH.LEFT
    if logo_stream is not None:
        p_logo.add_run().add_picture(logo_stream, width=Inches(2.15))
    elif pic_path is not None and pic_path.is_file():
        p_logo.add_run().add_picture(str(pic_path), width=Inches(2.15))
    else:
        r0 = p_logo.add_run("La Internacional\nde Seguros")
        _set_run_font(r0, bold=True, size=8, color=TITLE_BLUE)

    if pic_is_temp and pic_path is not None:
        try:
            pic_path.unlink(missing_ok=True)
        except OSError:
            pass

    pr = c_txt.paragraphs[0]
    pr.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    pr.paragraph_format.right_indent = Pt(8)
    r1 = pr.add_run("La Internacional de Seguros, C.A.\n")
    _set_run_font(r1, bold=True, size=7, color=TITLE_BLUE)
    r2 = pr.add_run("Plan de Optimización y Rentabilidad (POR) 2026\n")
    _set_run_font(r2, size=6)
    r3 = pr.add_run("Informe confidencial — Reaseguro y TPA en Salud")
    _set_run_font(r3, italic=True, size=6, color=ACCENT_GOLD)

    doc.add_paragraph()
    sep = doc.add_paragraph()
    run_sep = sep.add_run("―" * 62)
    _set_run_font(run_sep, size=8, color=ACCENT_GOLD)

    t = doc.add_paragraph()
    t.alignment = WD_ALIGN_PARAGRAPH.CENTER
    tr = t.add_run("INFORME CONFIDENCIAL\n")
    _set_run_font(tr, bold=True, size=15, color=TITLE_BLUE)
    tr2 = t.add_run("Análisis de reaseguro, TPA en salud y procedimiento administrativo (PAS)")
    _set_run_font(tr2, bold=True, size=10, color=TITLE_BLUE)

    doc.add_paragraph()
    minuta_act._add_table(
        doc,
        ["Campo", "Detalle"],
        [
            ["Referencia", referencia],
            ["Clasificación", "Confidencial — circulación restringida"],
            ["Fecha del documento", fecha_documento],
            ["Norma / contexto", "Solicitud consolidada v2; expediente SUDEASEG / SAA (reaseguro)"],
            [
                "Ubicación de respaldo",
                "Custodia y respaldo bajo **Gerencia Legal** y equipos de **apoyo técnico**; circulación acotada a **alta dirección**.",
            ],
            ["Elaborado por", elaborado_por],
        ],
    )


def build_confidential_document(
    *,
    out_path: Path,
    fecha_documento: str,
    referencia: str,
    elaborado_por: str,
    user_source_docx: Path | None,
    skip_user_body: bool,
    incluir_matriz_detallada: bool = False,
    fusion_completa: bool = False,
    ratio_resumen: float = 0.33,
) -> None:
    if fusion_completa:
        _build_confidential_document_fusion_legacy(
            out_path=out_path,
            fecha_documento=fecha_documento,
            referencia=referencia,
            elaborado_por=elaborado_por,
            user_source_docx=user_source_docx,
            skip_user_body=skip_user_body,
            incluir_matriz_detallada=incluir_matriz_detallada,
        )
        return

    doc = Document()
    _build_header_minuta_accionista(
        doc,
        fecha_documento=fecha_documento,
        referencia=referencia,
        elaborado_por=elaborado_por,
    )

    _append_propuesta_v1_bloques_datos(doc)
    _rule_gold(doc)

    _append_anexo_a_sin_remision_sudeaseg(doc)

    _rule_gold(doc)

    doc.add_paragraph()
    foot = doc.add_paragraph()
    fr = foot.add_run(
        f"Confidencial — La Internacional de Seguros — Uso restringido. Fecha: {fecha_documento}."
    )
    _set_run_font(fr, italic=True, size=8)

    out_path.parent.mkdir(parents=True, exist_ok=True)
    doc.save(str(out_path))


def _build_confidential_document_fusion_legacy(
    *,
    out_path: Path,
    fecha_documento: str,
    referencia: str,
    elaborado_por: str,
    user_source_docx: Path | None,
    skip_user_body: bool,
    incluir_matriz_detallada: bool,
) -> None:
    doc = Document()
    _build_header_confidencial(
        doc,
        fecha_documento=fecha_documento,
        referencia=referencia,
        elaborado_por=elaborado_por,
    )

    _sintesis_ejecutiva_portada(doc)
    _rule_gold(doc)

    if not skip_user_body and user_source_docx is not None:
        _append_user_analysis_integrated(doc, user_source_docx)
        _rule_gold(doc)

    minuta_act.append_confidential_reinsurance_tpa_sections(
        doc,
        para_accionista=True,
        incluir_matriz_detallada=incluir_matriz_detallada,
        modo_informe_ejecutivo=True,
    )

    doc.add_paragraph()
    foot = doc.add_paragraph()
    fr = foot.add_run(
        f"Confidencial — La Internacional de Seguros — Uso restringido. Fecha: {fecha_documento}."
    )
    _set_run_font(fr, italic=True, size=8)

    out_path.parent.mkdir(parents=True, exist_ok=True)
    doc.save(str(out_path))


def _default_user_downloads_path() -> Path | None:
    downloads = Path.home() / "Downloads"
    if not downloads.is_dir():
        return None
    for p in downloads.glob("*Reaseguro*TPA*.docx"):
        if p.name.startswith("~$"):
            continue
        return p
    for p in downloads.glob("*Reaseguro*.docx"):
        if p.name.startswith("~$"):
            continue
        if "TPA" in p.name.upper() or "salud" in p.name.lower():
            return p
    return None


def main() -> None:
    ap = argparse.ArgumentParser()
    ap.add_argument("--out", type=Path, default=DEFAULT_OUT)
    ap.add_argument("--fecha-doc", default="2026-04-24")
    ap.add_argument("--referencia", default="CONF-REA-TPA-SALUD-20260424-01")
    ap.add_argument("--elaborado-por", default="[Nombre y cargo]")
    ap.add_argument(
        "--desde-docx",
        type=Path,
        default=None,
        help="Ruta opcional a un Word de trabajo (no se inserta en el cuerpo de la minuta v1).",
    )
    ap.add_argument(
        "--copia-usuario",
        type=Path,
        default=None,
        help="Ruta opcional para guardar una segunda copia (p. ej. Descargas).",
    )
    ap.add_argument(
        "--sin-cuerpo-usuario",
        action="store_true",
        help="No fusionar el DOCX del usuario (solo bloque corporativo/técnico).",
    )
    ap.add_argument(
        "--matriz-detallada",
        action="store_true",
        help="Solo aplica con --fusion-completa: añade tabla SUDEASEG 3.2 bis.",
    )
    ap.add_argument(
        "--fusion-completa",
        action="store_true",
        help="Flujo antiguo: síntesis fija + fusión literal del DOCX + bloque técnico append.",
    )
    ap.add_argument(
        "--ratio-resumen",
        type=float,
        default=0.33,
        help="Proporción orientativa del texto fuente a conservar (predeterminado 0,33).",
    )
    args = ap.parse_args()

    user_path = args.desde_docx
    if user_path is None:
        user_path = _default_user_downloads_path()

    build_confidential_document(
        out_path=args.out,
        fecha_documento=args.fecha_doc,
        referencia=args.referencia,
        elaborado_por=args.elaborado_por,
        user_source_docx=user_path,
        skip_user_body=args.sin_cuerpo_usuario,
        incluir_matriz_detallada=args.matriz_detallada,
        fusion_completa=args.fusion_completa,
        ratio_resumen=args.ratio_resumen,
    )
    print(f"DOCX confidencial: {args.out}")
    if user_path:
        print(f"Fuente usuario fusionada: {user_path}")
    else:
        print("Aviso: no se encontró DOCX del usuario en Descargas; use --desde-docx.")

    if args.copia_usuario:
        args.copia_usuario.parent.mkdir(parents=True, exist_ok=True)
        shutil.copy2(args.out, args.copia_usuario)
        print(f"Copia adicional: {args.copia_usuario}")


if __name__ == "__main__":
    main()
