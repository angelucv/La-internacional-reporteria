"""
Genera INDICE-Y-AVANCE-Carpeta-Mesa-Actuarial.docx en Reportes/Mesa Actuarial.
Describe el propósito de cada subcarpeta y el avance general (POR 2026).
"""
from __future__ import annotations

from pathlib import Path

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.shared import Pt, RGBColor

_WORKSPACE_ROOT = Path(__file__).resolve().parents[2]
OUT_PATH = (
    _WORKSPACE_ROOT
    / "Reportes"
    / "Mesa Actuarial"
    / "INDICE-Y-AVANCE-Carpeta-Mesa-Actuarial.docx"
)

# Fecha de la última revisión sustantiva del índice (texto visible y pie).
FECHA_REVISION_INDICE = "2026-04-24"

TITLE_BLUE = RGBColor(0x1B, 0x3A, 0x5C)
BODY = RGBColor(0x33, 0x33, 0x33)


def _style_normal(doc: Document) -> None:
    n = doc.styles["Normal"]
    n.font.name = "Arial"
    n.font.size = Pt(11)
    n.font.color.rgb = BODY
    if n._element.rPr is not None and n._element.rPr.rFonts is not None:
        n._element.rPr.rFonts.set(qn("w:eastAsia"), "Arial")


def _h(doc: Document, text: str, level: int = 1) -> None:
    p = doc.add_heading(text, level=level)
    for r in p.runs:
        r.font.name = "Arial"
        r.font.color.rgb = TITLE_BLUE
        r.font.bold = True
        r.font.size = Pt(16 if level == 1 else 13)


def _p(doc: Document, text: str) -> None:
    para = doc.add_paragraph()
    run = para.add_run(text)
    run.font.name = "Arial"
    run.font.size = Pt(11)
    run.font.color.rgb = BODY
    para.paragraph_format.space_after = Pt(8)


def _bullet(doc: Document, text: str) -> None:
    para = doc.add_paragraph(style="List Bullet")
    run = para.add_run(text)
    run.font.name = "Arial"
    run.font.size = Pt(11)
    run.font.color.rgb = BODY


def build() -> None:
    doc = Document()
    _style_normal(doc)

    t = doc.add_paragraph()
    tr = t.add_run("Índice y avance de la carpeta\nMesa Actuarial — Plan de Optimización y Rentabilidad 2026")
    tr.bold = True
    tr.font.size = Pt(18)
    tr.font.color.rgb = TITLE_BLUE
    tr.font.name = "Arial"
    t.alignment = WD_ALIGN_PARAGRAPH.CENTER
    doc.add_paragraph()

    sub = doc.add_paragraph()
    sr = sub.add_run(
        "Documento orientado a quien consulte el contenido de la carpeta de trabajo de la mesa actuarial "
        "(Drive u otro repositorio compartido). Resume el propósito de cada subcarpeta y el estado de avance "
        "a la fecha de elaboración."
    )
    sr.font.name = "Arial"
    sr.font.size = Pt(11)
    sr.italic = True
    sr.font.color.rgb = BODY
    sub.alignment = WD_ALIGN_PARAGRAPH.CENTER
    doc.add_paragraph()

    _h(doc, "1. Propósito de esta carpeta")
    _p(
        doc,
        "Concentra el material que la mesa actuarial va produciendo o recibiendo en el marco del "
        "Plan de Optimización y Rentabilidad 2026: comunicaciones oficiales, reportería de gestión, "
        "soportes de producto, minutas de reunión, insumos para el tablero de información financiera-sectorial "
        "y el seguimiento semanal del trabajo. No sustituye los sistemas corporativos ni el expediente regulatorio; "
        "sirve de punto único para coordinación, asesoría y gobierno del proyecto.",
    )

    _h(doc, "2. Avance general (visión sintética)")
    _p(
        doc,
        "La mesa actuarial dispone ya de un **corpus amplio** promovido desde la carpeta compartida de gerencia "
        "actuarial hacia esta carpeta de trabajo (espejo bajo «Materiales», 176 archivos a 2026-04-23), con "
        "**traza explícita** en Markdown y minuta de validación frente a la Solicitud consolidada v2 "
        "(código POR de la minuta vigente en `Minutas Reuniones`, p. ej. MIN-POR-ACT-20260424-04). El tablero de datos sectorial y la reportería de gestión por cierres "
        "siguen alineados al uso en asesoría y en el agente compilador del POR.",
    )
    _p(
        doc,
        "Los **huecos prioritarios** frente al pliego siguen siendo: **SEFA** y comprobantes analíticos, "
        "**volcado transaccional de primas**, **profundización contractual de reaseguro** y homologación fina "
        "con plantillas **SUDEASEG** en siniestros. El objetivo de la semana en curso es pasar de «tener "
        "insumos» a «cerrar lectura» para que el compilador multi-mesa y el Drive reciban un paquete "
        "**homogéneo y citado**.",
    )
    _p(
        doc,
        "El **control semanal** se actualiza en la subcarpeta «Semanal»: existen borradores Semana 1 y, a partir "
        f"de {FECHA_REVISION_INDICE}, la **Semana 2** en PowerPoint (`La_Internacional_Semana2_2026.pptx` y "
        "`REP-POR-ACT-Semana02-2026-BORRADOR.pptx`) con una lámina de síntesis de avances en la diapositiva de "
        "cierre semanal.",
    )

    _h(doc, "2.1 Avance de esta semana (Semana 2 — para el agente y el equipo)")
    _bullet(
        doc,
        "Insumos actuariales **centralizados** en `Materiales/Promocion-2026-04-23-Carpeta-compartida-Actuarial` "
        "(baremos 2026, margen de solvencia 2025, gestión mensual, IBNR, RiesgCAT, reaseguro SUDEASEG, "
        "certificaciones de reservas, siniestros cierre 2025, etc.).",
    )
    _bullet(
        doc,
        "Minuta con **fecha de actualización en el nombre de archivo**: ver `Minutas Reuniones` "
        "(patrón `MIN-POR-ACT-AAAAMMDD-Mesa-Actuarial-vs-Solicitud_actualizado-AAAA-MM-DD.docx`).",
    )
    _bullet(
        doc,
        "Presentación Semana 2: **diapositiva de control semanal** ampliada con cinco viñetas de avance "
        "(promoción, minuta, índice, pendientes pliego, siguiente paso Drive/compilador).",
    )
    _bullet(
        doc,
        "Próximo foco operativo: cerrar **SEFA / analítica**, **primas**, **reaseguro contractual** y "
        "referenciar acuerdos en `la-internacional-reporteria` cuando haya decisión formal.",
    )
    _bullet(
        doc,
        "Nuevas carpetas bajo `Reportes/` desde **Gerencia Legal**: "
        "`CONDICIONADOS DE POLIZAS - PRODUCTOS`, "
        "`CONTRATOS DE REASEGUROS-20260423T200206Z-3-001` y el árbol `GERENCIA LEGAL/` "
        "(contratos, intermediarios, procedimientos sancionatorios, juicios, etc.).",
    )
    _bullet(
        doc,
        "Matriz de trabajo **reaseguro Legal vs SUDEASEG**: "
        "`Materiales/MATRIZ-REASEGURO-Legal-vs-SUDEASEG-2026-04-24.md` (listado de PDF y tabla para completar el cruce).",
    )

    _h(doc, "3. Contenido por subcarpeta")

    _h(doc, "3.1 BI financiero", level=2)
    _p(
        doc,
        "Reúne los conjuntos de datos que alimentan la aplicación web de información sectorial y financiera "
        "(series de primas, índices del boletín, tipo de cambio de referencia, indicadores anuarios y cadena de "
        "resultado técnico cuando aplica), más notas de contexto sobre el desarrollo del tablero y del demo "
        "asociado. Su función es permitir lectura paralela entre lo público-consolidado y el trabajo interno "
        "de la mesa.",
    )
    _bullet(doc, "Avance: conjunto de datos completo para el tablero en su versión actual; documentación de apoyo disponible.")

    _h(doc, "3.2 Correos", level=2)
    _p(
        doc,
        "Almacena mensajes institucionales en soporte fijo (habitualmente PDF) relativos a cierres de gestión "
        "y al detalle de productos remitidos en la misma ventana temporal. Sirven de prueba de remisión y "
        "contexto frente a la mesa.",
    )
    _bullet(doc, "Avance: tres envíos de referencia ya archivados para el periodo en curso.")

    _h(doc, "3.3 Materiales", level=2)
    _p(
        doc,
        "Contiene documentación narrativa de la oferta (productos y servicios por ramo), elaborada y revisada "
        "según la cadena interna de aprobación. Complementa lo conversado en sesión y lo recogido por correo.",
    )
    _bullet(doc, "Avance: material maestro de producto disponible para consulta de la mesa y de asesoría.")

    _h(doc, "3.4 Minutas Reuniones", level=2)
    _p(
        doc,
        "Guarda las actas de las sesiones de la mesa actuarial bajo la codificación del plan POR, con fecha "
        "de reunión y correlativo. Son la fuente de acuerdos y del calendario de seguimiento.",
    )
    _bullet(
        doc,
        "Avance: sesiones de abril registradas; además, minuta de validación carpeta vs Solicitud v2 con "
        "nombre que incluye fecha de actualización y código POR correlativo para trazabilidad ante compilador.",
    )

    _h(doc, "3.5 Reportes", level=2)
    _p(
        doc,
        "Agrupa la reportería de gestión interna por cierres relevantes: paquetes con visión de gestión general, "
        "inventarios, siniestralidad por ramos destacados (salud, automóvil) e intermediarios, además de un resumen "
        "auxiliar vinculado al programa de reaseguro para lectura conjunta con la documentación contractual cuando "
        "ésta se incorpore.",
    )
    _bullet(doc, "Avance: dos paquetes de cierre ya ubicados aquí; revisión técnica en curso frente al pliego del diagnóstico.")

    _h(doc, "3.6 Semanal", level=2)
    _p(
        doc,
        "Destinada al seguimiento por ciclo corto: borrador de informe de gestión, hallazgos y controles para la "
        "primera oleada del plan, en formatos de texto y presentación, con tono informativo para el control "
        "del trabajo de la mesa y como insumo para otros foros de seguimiento si se requiere.",
    )
    _bullet(
        doc,
        "Avance: Semana 1 archivada (PDF y borrador PPTX/Markdown); **Semana 2** añadida con "
        "`La_Internacional_Semana2_2026.pptx` y `REP-POR-ACT-Semana02-2026-BORRADOR.pptx` (script "
        "`crear_presentacion_semana2_mesa_actuarial.py` en reportería).",
    )

    _h(doc, "3.7 Insumos promovidos desde Carpeta compartida Actuarial", level=2)
    _p(
        doc,
        "Subcarpeta bajo «Materiales» que concentra una copia íntegra (espejo) de lo recibido por la gerencia actuarial "
        "en «Reportes/Carpeta compartida Actuarial», promovida el 2026-04-23 a «Mesa Actuarial» para alimentar el "
        "Drive y el agente compilador del POR sin alterar el original en la bandeja de entrada. Incluye baremos "
        "convenidos 2026, margen de solvencia 2025, reportería de gestión mensual y el resto de archivos de primer "
        "nivel (176 archivos en total en la promoción registrada).",
    )
    _bullet(
        doc,
        "Ruta: Materiales/Promocion-2026-04-23-Carpeta-compartida-Actuarial. Trazabilidad en Markdown: "
        "INDICE-PROMOCION-CARPETA-ACTUARIAL-A-MESA.md en la raíz de esta carpeta.",
    )

    _h(doc, "4. Mantenimiento de este índice")
    _p(
        doc,
        "Conviene actualizar este documento cuando se creen subcarpetas nuevas, se cierre una oleada del plan "
        "o cambie sustancialmente el estado de avance. La fecha de la última revisión puede consignarse al pie.",
    )

    foot = doc.add_paragraph()
    fr = foot.add_run(
        "Clasificación: uso interno / confidencial según política de la compañía. "
        f"Última revisión de contenido: {FECHA_REVISION_INDICE}. "
        "Generado desde el repositorio de reportería (script generar_indice_carpeta_mesa_actuarial_docx.py)."
    )
    fr.font.name = "Arial"
    fr.font.size = Pt(9)
    fr.italic = True
    fr.font.color.rgb = BODY

    OUT_PATH.parent.mkdir(parents=True, exist_ok=True)
    try:
        doc.save(OUT_PATH)
        print(OUT_PATH)
    except PermissionError:
        alt = OUT_PATH.with_name(
            f"{OUT_PATH.stem}-REGENERADO-{FECHA_REVISION_INDICE}{OUT_PATH.suffix}"
        )
        doc.save(alt)
        print(alt)
        print(
            "(AVISO: no se pudo sobrescribir el archivo principal porque está abierto en otro programa; "
            "ciérrelo y vuelva a ejecutar el script, o renombre/reemplace manualmente.)",
        )


if __name__ == "__main__":
    build()
