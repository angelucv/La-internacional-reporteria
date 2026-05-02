"""
Genera minuta/informe en Word sobre la carga documental en la carpeta
compartida Mesa Actuarial y constancia de reunión con la mesa actuarial.

Logo: `plantillas/logo-la-internacional.png` (horizontal con símbolo). Respaldo: `Logo horizontal.svg`.
"""
from __future__ import annotations

import argparse
import io
import sys
import tempfile
from pathlib import Path

from docx import Document
from docx.enum.table import WD_ALIGN_VERTICAL, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor

TITLE_BLUE = RGBColor(0x1B, 0x3A, 0x5C)
ACCENT_GOLD = RGBColor(0xC9, 0xA2, 0x27)

ROOT_REPORTERIA = Path(__file__).resolve().parents[1]
WORKSPACE = ROOT_REPORTERIA.parent
DEFAULT_LOGO_SVG = WORKSPACE / "La-Internacional-BI" / "public" / "Images" / "Logo horizontal.svg"
BUNDLED_LOGO_PNG = ROOT_REPORTERIA / "plantillas" / "logo-la-internacional.png"
DEFAULT_OUT = (
    WORKSPACE
    / "Reportes"
    / "Mesa Actuarial"
    / "Minutas Reuniones"
    / "MIN-POR-ACT-20260422-Informe-Carpeta-Actuarial-Yartiza-Oberto.docx"
)


def _set_run_font(run, *, bold: bool = False, italic: bool = False, size: int = 11, color=None) -> None:
    run.bold = bold
    run.italic = italic
    run.font.name = "Arial"
    run.font.size = Pt(size)
    run._element.rPr.rFonts.set(qn("w:eastAsia"), "Arial")
    if color is not None:
        run.font.color.rgb = color


def _heading(doc: Document, text: str, level: int = 1) -> None:
    h = doc.add_heading(text, level=level)
    sec_size = 14 if level == 1 else 11
    for run in h.runs:
        _set_run_font(run, bold=True, size=sec_size)
        run.font.color.rgb = TITLE_BLUE
    h.paragraph_format.space_before = Pt(12)
    h.paragraph_format.space_after = Pt(6)


def _p(doc: Document, text: str, *, bold: bool = False, italic: bool = False, size: int = 11) -> None:
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(6)
    r = p.add_run(text)
    _set_run_font(r, bold=bold, italic=italic, size=size)


def _add_table(doc: Document, headers: list[str], rows: list[list[str]]) -> None:
    table = doc.add_table(rows=1 + len(rows), cols=len(headers))
    table.style = "Table Grid"
    for j, h in enumerate(headers):
        cell = table.rows[0].cells[j]
        cell.text = ""
        p = cell.paragraphs[0]
        run = p.add_run(h)
        _set_run_font(run, bold=True, size=10)
        p.paragraph_format.space_after = Pt(3)
    for i, row in enumerate(rows, start=1):
        for j, val in enumerate(row):
            cell = table.rows[i].cells[j]
            cell.text = ""
            p = cell.paragraphs[0]
            run = p.add_run(val)
            _set_run_font(run, size=10)
    doc.add_paragraph()


def _svg_to_png(svg_path: Path) -> Path | None:
    try:
        from reportlab.graphics import renderPM
        from svglib.svglib import svg2rlg
    except ImportError:
        return None
    if not svg_path.is_file():
        return None
    drawing = svg2rlg(str(svg_path))
    if drawing is None:
        return None
    tmp = Path(tempfile.mkstemp(suffix=".png")[1])
    renderPM.drawToFile(drawing, str(tmp), fmt="PNG", dpi=144)
    return tmp


def _resolve_logo_path(*, logo_svg: Path, logo_png: Path | None) -> tuple[Path | None, bool]:
    if logo_png is not None and logo_png.is_file():
        return logo_png, False
    if BUNDLED_LOGO_PNG.is_file():
        return BUNDLED_LOGO_PNG, False
    tmp = _svg_to_png(logo_svg)
    if tmp is not None and tmp.is_file():
        return tmp, True
    return None, False


def _logo_png_bytesio(pic_path: Path) -> io.BytesIO | None:
    try:
        from PIL import Image
    except ImportError:
        return None
    try:
        im = Image.open(pic_path)
        if im.mode not in ("RGB", "L"):
            im = im.convert("RGB")
        buf = io.BytesIO()
        im.save(buf, format="PNG", optimize=True)
        buf.seek(0)
        return buf
    except OSError:
        return None


def build_document(
    *,
    out_path: Path,
    logo_svg: Path,
    logo_png: Path | None,
    fecha_documento: str,
    fecha_reunion_yartiza: str,
    elaborado_por: str,
) -> None:
    doc = Document()
    style = doc.styles["Normal"]
    style.font.name = "Arial"
    style.font.size = Pt(11)
    style._element.rPr.rFonts.set(qn("w:eastAsia"), "Arial")

    pic_path, pic_is_temp = _resolve_logo_path(logo_svg=logo_svg, logo_png=logo_png)
    logo_stream = _logo_png_bytesio(pic_path) if pic_path is not None else None

    enc = doc.add_table(rows=1, cols=2)
    enc.autofit = False
    enc.alignment = WD_TABLE_ALIGNMENT.CENTER
    w_logo_col = Inches(3.05)
    w_txt_col = Inches(3.15)
    enc.columns[0].width = w_logo_col
    enc.columns[1].width = w_txt_col
    c_logo = enc.rows[0].cells[0]
    c_txt = enc.rows[0].cells[1]
    c_logo.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
    c_txt.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
    c_logo.width = w_logo_col
    c_txt.width = w_txt_col

    p_logo = c_logo.paragraphs[0]
    p_logo.alignment = WD_ALIGN_PARAGRAPH.LEFT
    p_logo.paragraph_format.space_before = Pt(0)
    p_logo.paragraph_format.space_after = Pt(0)
    if logo_stream is not None:
        p_logo.add_run().add_picture(logo_stream, width=Inches(2.75))
    elif pic_path is not None and pic_path.is_file():
        p_logo.add_run().add_picture(str(pic_path), width=Inches(2.75))
    else:
        r = p_logo.add_run("La Internacional\nde Seguros")
        _set_run_font(r, bold=True, size=8, color=TITLE_BLUE)

    if pic_is_temp and pic_path is not None:
        try:
            pic_path.unlink(missing_ok=True)
        except OSError:
            pass

    pr = c_txt.paragraphs[0]
    pr.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    pr.paragraph_format.space_before = Pt(0)
    pr.paragraph_format.space_after = Pt(0)
    pr.paragraph_format.right_indent = Pt(8)
    r1 = pr.add_run("La Internacional de Seguros, C.A.\n")
    _set_run_font(r1, bold=True, size=9, color=TITLE_BLUE)
    r2 = pr.add_run("Plan de Optimización y Rentabilidad (POR) 2026\n")
    _set_run_font(r2, size=8)
    r3 = pr.add_run("Mesa Actuarial · Reportería documental")
    _set_run_font(r3, italic=True, size=8, color=ACCENT_GOLD)

    doc.add_paragraph()
    sep = doc.add_paragraph()
    sep.paragraph_format.space_after = Pt(2)
    run_sep = sep.add_run("―" * 62)
    _set_run_font(run_sep, size=8, color=ACCENT_GOLD)

    t = doc.add_paragraph()
    t.alignment = WD_ALIGN_PARAGRAPH.CENTER
    tr = t.add_run("MINUTA — INFORME DE CARGA DOCUMENTAL\n")
    _set_run_font(tr, bold=True, size=18, color=TITLE_BLUE)
    tr2 = t.add_run("Carpeta compartida Mesa Actuarial")
    _set_run_font(tr2, bold=True, size=12, color=TITLE_BLUE)

    doc.add_paragraph()

    _add_table(
        doc,
        ["Campo", "Detalle"],
        [
            ["Referencia sugerida", "MIN-POR-ACT-20260422-01"],
            ["Fecha del presente informe", fecha_documento],
            ["Asunto", "Constancia de reunión y detalle del material disponible en la carpeta compartida"],
            ["Ámbito", "Reportes\\Mesa Actuarial (sincronizada con Drive / uso interno POR)"],
            ["Elaborado por", elaborado_por],
        ],
    )

    _heading(doc, "1. Documentación de la reunión con la Lic. Yartiza Oberto", level=1)
    _p(
        doc,
        f"En fecha {fecha_reunion_yartiza} se sostuvo una reunión con la Lic. Yartiza Oberto, "
        "quien informó que se había completado la carga de la información actuarial y de apoyo "
        "en la carpeta compartida de la Mesa Actuarial, a fin de disponer de los insumos para "
        "seguimiento del POR, reportería hacia accionistas y alimentación de los tableros de BI "
        "cuando corresponda.",
    )
    _p(
        doc,
        "De manera resumida, lo comunicado en la sesión se relaciona con: (i) disponibilidad de "
        "correos y anexos en PDF que contextualizan los cierres de gestión; (ii) paquetes de "
        "reportes de gestión en Excel para los cortes diciembre 2025 y marzo 2026; (iii) material "
        "de productos y servicios; (iv) conjunto de archivos CSV para el tablero financiero; "
        "(v) revisiones documentadas del BI web y del demo; y (vi) borradores de reportería semanal.",
        italic=False,
    )
    _p(
        doc,
        "Los nombres exactos de archivos y rutas relativas figuran en el inventario del apartado 2, "
        "para trazabilidad y verificación en campo.",
        italic=True,
    )

    _heading(doc, "2. Inventario de la información cargada en la carpeta compartida", level=1)
    _p(
        doc,
        "A continuación se detalla el contenido verificado en el workspace local "
        r"(…\Reportes\Mesa Actuarial), coherente con la estructura prevista para la carpeta compartida. "
        "Las fechas de modificación registradas en disco para los archivos listados fueron predominantemente el 15-04-2026.",
    )

    _heading(doc, "2.1. Raíz de la carpeta", level=2)
    _add_table(
        doc,
        ["Archivo", "Descripción breve"],
        [
            [
                "INDICE-Y-AVANCE-Carpeta-Mesa-Actuarial.docx",
                "Índice y avance de la carpeta (documento de referencia interna).",
            ],
        ],
    )

    _heading(doc, "2.2. Subcarpeta «BI financiero»", level=2)
    _p(doc, "Archivos tabulares (CSV) para tableros y análisis, más revisiones en Word.")
    _add_table(
        doc,
        ["Archivo", "Contenido / uso"],
        [
            ["bcv_ves_por_usd_mensual.csv", "Serie mensual tipo de cambio referencial (VES/USD)."],
            ["cuadro_29_indicadores_financieros_2023_por_empresa.csv", "Indicadores financieros 2023 por empresa."],
            [
                "cuadro_31A_primas_netas_cobradas_2023_vs_2022.csv",
                "Primas netas cobradas 2023 frente a 2022 (cuadro 31A).",
            ],
            ["indices_por_empresa_historico_largo.csv", "Índices por empresa — serie histórica extendida."],
            ["indices_por_empresa_mes_actual.csv", "Índices por empresa — corte mes en curso."],
            ["primas_netas_mensual_largo.csv", "Primas netas mensuales — serie larga."],
            ["resultado_tecnico_saldo_mensual.csv", "Resultado técnico / saldos mensuales."],
            ["Revision-1-BI-Web-La-Internacional.docx", "Revisión 1 del BI web La Internacional."],
            ["Revision-2-Demo-Datos-Public-y-Drive.docx", "Revisión 2 del demo de datos públicos y Drive."],
        ],
    )

    _heading(doc, "2.3. Subcarpeta «Correos»", level=2)
    _p(doc, "Correos de referencia en PDF (remisión 13-04-2026).")
    _add_table(
        doc,
        ["Archivo", "Tema"],
        [
            ["Correo 1 13-04-2026 cierre marzo 2026.pdf", "Cierre y reportes de gestión — marzo 2026."],
            ["Correo 2 13-04-2026 cierre diciembre 2025.pdf", "Cierre y reportes de gestión — diciembre 2025."],
            [
                "Correo 3 13-04-2026 contenido productos la internacional.pdf",
                "Contenido de productos La Internacional.",
            ],
        ],
    )

    _heading(doc, "2.4. Subcarpeta «Materiales»", level=2)
    _add_table(
        doc,
        ["Archivo", "Descripción"],
        [
            [
                "Contenido de Productos y Servicios La Internacional de Seguros.docx",
                "Documento corporativo de catálogo / descripción de oferta.",
            ],
        ],
    )

    _heading(doc, "2.5. Subcarpeta «Minutas Reuniones»", level=2)
    _add_table(
        doc,
        ["Archivo", "Descripción"],
        [
            ["MIN-POR-ACT-20260413-01.docx", "Minuta de reunión 13-04-2026."],
            ["MIN-POR-ACT-20260414-01.docx", "Minuta de reunión 14-04-2026."],
        ],
    )

    _heading(doc, "2.6. Subcarpeta «Reportes»", level=2)
    _p(doc, "Paquetes de gestión por corte y archivo complementario de reaseguros.")
    _add_table(
        doc,
        ["Ruta / archivo", "Descripción"],
        [
            [
                r"reportesdegestiongeneraldiciembre2025\GESTION GENERAL DICIEMBRE 2025.xlsx",
                "Gestión general al cierre diciembre 2025.",
            ],
            [
                r"reportesdegestiongeneraldiciembre2025\INVENTARIO DE POLIZAS AÑO 2025.xlsx",
                "Inventario de pólizas — año 2025.",
            ],
            [
                r"reportesdegestiongeneraldiciembre2025\ESTADISTICAS DE AUTOMOVIL 2025.xlsx",
                "Estadísticos automóvil 2025 (libro de gran tamaño).",
            ],
            [
                r"reportesdegestiongeneraldiciembre2025\REPORTE DE SINIESTROS SALUD DICIEMBRE.xlsx",
                "Siniestros salud — diciembre (dato sensible; manejo acotado).",
            ],
            [
                r"reportesdegestiongeneraldiciembre2025\SINIESTRALIDAD INTERMEDIARIOS DICIEMBRE (AÑO 2025) POR RAMO.xlsx",
                "Siniestralidad por intermediarios — diciembre 2025 por ramo.",
            ],
            [
                r"reportesdegestionmensualalcierredelmesdemarzo2026\GESTION GENERAL MARZO 2026.xlsx",
                "Gestión general al cierre marzo 2026.",
            ],
            [
                r"reportesdegestionmensualalcierredelmesdemarzo2026\INVENTARIO DE POLIZAS MARZO 2026.xlsx",
                "Inventario de pólizas — marzo 2026.",
            ],
            [
                r"reportesdegestionmensualalcierredelmesdemarzo2026\REPORTE ESTADISTICO DE AUTOMOVIL ALCIERRE DE MARZO DEL 2026.xlsx",
                "Estadísticos automóvil al cierre marzo 2026 (libro de gran tamaño).",
            ],
            [
                r"reportesdegestionmensualalcierredelmesdemarzo2026\REPORTE DE SINIESTROS SALUD MARZO US$.xlsx",
                "Siniestros salud — marzo en US$ (dato sensible; manejo acotado).",
            ],
            [
                r"reportesdegestionmensualalcierredelmesdemarzo2026\SINIESTRALIDAD INTERMEDIARIOS MARZO 2026 POR RAMO.xlsx",
                "Siniestralidad por intermediarios — marzo 2026 por ramo.",
            ],
            [
                "Resumen de pagos Reaseguros año 2025 Juan.xlsx",
                "Resumen de pagos de reaseguros (referencia 2025).",
            ],
        ],
    )

    _heading(doc, "2.7. Subcarpeta «Semanal»", level=2)
    _add_table(
        doc,
        ["Archivo", "Descripción"],
        [
            ["REP-POR-ACT-Semana01-2026-BORRADOR.md", "Borrador Markdown — reporte semanal mesa actuarial."],
            ["REP-POR-ACT-Semana01-2026-BORRADOR.pptx", "Presentación asociada al borrador semanal."],
        ],
    )

    _heading(doc, "3. Síntesis para la reunión con el liderazgo del proyecto", level=1)
    _p(
        doc,
        "La carpeta compartida Mesa Actuarial concentra, en una estructura ordenada por temas, "
        "los insumos para seguimiento actuarial y de gestión: cierres 2025-12 y 2026-03, correos "
        "de contexto, material de productos, datos para BI financiero, revisiones del tablero web "
        "y del demo, así como borradores de reportería semanal. El canal está alineado con la "
        "política de reportería oficial hacia accionistas descrita en el documento maestro del workspace.",
    )

    _heading(doc, "4. Acuerdos y próximos pasos sugeridos", level=1)
    for text in [
        "Validar con el liderazgo del proyecto la prioridad de consumo de cada paquete (BI interno vs. tableros públicos).",
        "Mantener el índice INDICE-Y-AVANCE-Carpeta-Mesa-Actuarial.docx actualizado cuando ingresen nuevos cortes.",
        "Respetar políticas de datos sensibles en salud y automóvil antes de cualquier despliegue fuera de red controlada.",
        "Registrar en el repositorio de reportería (minutas POR) cualquier decisión formal derivada de la sesión de liderazgo.",
    ]:
        p = doc.add_paragraph(style="List Number")
        r = p.add_run(text)
        _set_run_font(r, size=11)

    doc.add_paragraph()
    foot = doc.add_paragraph()
    fr = foot.add_run(
        f"Documento generado para fines de trabajo POR. Fecha de redacción: {fecha_documento}. "
        "Las manifestaciones sobre el contenido de la reunión con la Lic. Yartiza Oberto deben "
        "ajustarse, si fuere necesario, a la versión aprobada por las partes."
    )
    _set_run_font(fr, italic=True, size=9)

    out_path.parent.mkdir(parents=True, exist_ok=True)
    doc.save(str(out_path))


def main() -> None:
    ap = argparse.ArgumentParser()
    ap.add_argument("--out", type=Path, default=DEFAULT_OUT)
    ap.add_argument("--logo-svg", type=Path, default=DEFAULT_LOGO_SVG)
    ap.add_argument("--logo-png", type=Path, default=None)
    ap.add_argument("--fecha-doc", default="2026-04-22")
    ap.add_argument("--fecha-reunion", default="2026-04-21")
    ap.add_argument("--elaborado-por", default="[Nombre y cargo]")
    args = ap.parse_args()
    build_document(
        out_path=args.out,
        logo_svg=args.logo_svg,
        logo_png=args.logo_png,
        fecha_documento=args.fecha_doc,
        fecha_reunion_yartiza=args.fecha_reunion,
        elaborado_por=args.elaborado_por,
    )
    print(f"DOCX generado: {args.out}")


if __name__ == "__main__":
    main()
