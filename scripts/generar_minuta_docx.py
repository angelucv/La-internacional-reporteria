"""
Genera .docx de la minuta (Arial, estructura alineada al markdown).
"""
from __future__ import annotations

import re
import sys
from pathlib import Path

from docx import Document
from docx.oxml.ns import qn
from docx.shared import Pt, RGBColor

ROOT = Path(__file__).resolve().parents[1]
DEFAULT_MD = ROOT / "reuniones" / "2026" / "MIN-POR-ACT-20260414-01.md"
DEFAULT_OUT = ROOT / "reuniones" / "2026" / "MIN-POR-ACT-20260414-01.docx"

TITLE_BLUE = RGBColor(0x1B, 0x3A, 0x5C)


def _set_run_font(run, *, bold: bool = False, italic: bool = False, size: int = 11) -> None:
    run.bold = bold
    run.italic = italic
    run.font.name = "Arial"
    run.font.size = Pt(size)
    run._element.rPr.rFonts.set(qn("w:eastAsia"), "Arial")


def _add_paragraph_with_inline(doc: Document, text: str, *, size: int = 11) -> None:
    """Interpreta **negrita** en una línea."""
    text = text.strip()
    if not text:
        doc.add_paragraph()
        return
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(6)
    parts = re.split(r"(\*\*.+?\*\*)", text)
    for part in parts:
        if not part:
            continue
        if part.startswith("**") and part.endswith("**"):
            run = p.add_run(part[2:-2])
            _set_run_font(run, bold=True, size=size)
        else:
            run = p.add_run(part)
            _set_run_font(run, size=size)


def _flush_table(doc: Document, rows: list[list[str]]) -> None:
    if not rows:
        return
    ncols = max(len(r) for r in rows)
    table = doc.add_table(rows=len(rows), cols=ncols)
    table.style = "Table Grid"
    for ri, row in enumerate(rows):
        for ci in range(ncols):
            cell = table.rows[ri].cells[ci]
            cell.text = ""
            val = row[ci] if ci < len(row) else ""
            val = re.sub(r"\*\*(.+?)\*\*", r"\1", val)
            p = cell.paragraphs[0]
            run = p.add_run(val)
            _set_run_font(run, bold=(ri == 0), size=10)
    doc.add_paragraph()


def build_docx(md_path: Path, out_path: Path) -> None:
    text = md_path.read_text(encoding="utf-8")
    lines = text.splitlines()
    doc = Document()
    # Estilo base
    style = doc.styles["Normal"]
    style.font.name = "Arial"
    style.font.size = Pt(11)
    style._element.rPr.rFonts.set(qn("w:eastAsia"), "Arial")

    i = 0
    table_rows: list[list[str]] = []
    in_table = False

    while i < len(lines):
        line = lines[i]

        if line.strip().startswith("|") and "|" in line[1:]:
            if re.match(r"^\|[\s\-:|]+\|$", line.strip()):
                i += 1
                continue
            cells = [c.strip() for c in line.strip().strip("|").split("|")]
            in_table = True
            table_rows.append(cells)
            i += 1
            continue

        if in_table:
            _flush_table(doc, table_rows)
            table_rows = []
            in_table = False

        if line.startswith("# "):
            h = doc.add_heading(line[2:].strip(), level=1)
            for run in h.runs:
                _set_run_font(run, bold=True, size=18)
            if h.runs:
                h.runs[0].font.color.rgb = TITLE_BLUE
            i += 1
            continue

        if line.startswith("## "):
            h = doc.add_heading(line[3:].strip(), level=2)
            for run in h.runs:
                _set_run_font(run, bold=True, size=12)
            if h.runs:
                h.runs[0].font.color.rgb = TITLE_BLUE
            i += 1
            continue

        if line.strip() == "---":
            i += 1
            continue

        if line.strip().startswith("- "):
            p = doc.add_paragraph(style="List Bullet")
            _add_paragraph_with_inline_to_para(p, line.strip()[2:])
            i += 1
            continue

        if re.match(r"^\d+\.\s", line.strip()):
            p = doc.add_paragraph(style="List Number")
            _add_paragraph_with_inline_to_para(p, re.sub(r"^\d+\.\s*", "", line.strip()))
            i += 1
            continue

        if line.strip().startswith("*") and line.strip().endswith("*"):
            _add_paragraph_with_inline(doc, line.strip().strip("*"))
            i += 1
            continue

        if not line.strip():
            i += 1
            continue

        _add_paragraph_with_inline(doc, line)
        i += 1

    if in_table:
        _flush_table(doc, table_rows)

    out_path.parent.mkdir(parents=True, exist_ok=True)
    doc.save(str(out_path))


def _add_paragraph_with_inline_to_para(p, text: str) -> None:
    text = text.strip()
    parts = re.split(r"(\*\*.+?\*\*)", text)
    for part in parts:
        if not part:
            continue
        if part.startswith("**") and part.endswith("**"):
            run = p.add_run(part[2:-2])
            _set_run_font(run, bold=True, size=11)
        else:
            run = p.add_run(part)
            _set_run_font(run, size=11)


def main() -> None:
    md = Path(sys.argv[1]) if len(sys.argv) > 1 else DEFAULT_MD
    out = Path(sys.argv[2]) if len(sys.argv) > 2 else DEFAULT_OUT
    if not md.is_file():
        print(f"No existe: {md}", file=sys.stderr)
        sys.exit(1)
    build_docx(md, out)
    print(f"DOCX generado: {out}")


if __name__ == "__main__":
    main()
